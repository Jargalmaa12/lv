import contextlib
import io
import json
import math
import os
import re
import tempfile
import time
import zipfile
from dataclasses import dataclass
from typing import Any
from xml.etree import ElementTree as ET

import requests
import streamlit as st
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

BUILD_ID = "ai-app-2026-02-08.4"


# ===================================================
# Small utils
# ===================================================
def clean_text(s: str) -> str:
    s = (s or "").replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def normalize_key(s: str) -> str:
    return clean_text(s).lower()


Q_PREFIX_RE = re.compile(r"^\s*(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*", re.IGNORECASE)


def strip_q_prefix(s: str) -> str:
    return clean_text(Q_PREFIX_RE.sub("", s or "", count=1))


ANSWER_GUIDE_INLINE_RE = re.compile(
    r"\bAnswer\s+(?:may|must|need(?:s)?)\s+(?:to\s+)?address\b[:\s-]*",
    re.IGNORECASE,
)
ANSWER_GUIDE_START_RE = re.compile(
    r"^\s*Answer\s+(?:may|must|need(?:s)?)\s+(?:to\s+)?address\b",
    re.IGNORECASE,
)
ANSWER_GUIDE_ANY_RE = re.compile(
    r"\bAnswer\s+(?:may|must|need(?:s)?)\s+(?:to\s+)?address\b",
    re.IGNORECASE,
)


def strip_answer_guide(text: str) -> str:
    t = clean_text(text)
    if not t:
        return ""
    m = ANSWER_GUIDE_INLINE_RE.search(t)
    if not m:
        return t
    return clean_text(t[: m.start()])


def trim_after_question_mark(text: str) -> str:
    t = clean_text(text)
    if not t:
        return ""
    if "?" not in t:
        return t
    qpos = t.find("?")
    return clean_text(t[: qpos + 1])


def trim_after_sentence_if_long(text: str, max_chars: int = 220) -> str:
    t = clean_text(text)
    if len(t) <= max_chars:
        return t
    # If assessor guide/sample answer is appended without '?', cut at the first sentence end.
    for sep in [". ", "; ", " - "]:
        pos = t.find(sep)
        if 20 <= pos <= max_chars:
            return clean_text(t[: pos + (1 if sep.startswith(".") else 0)])
    return clean_text(t[:max_chars])


def is_red_hex(val: str) -> bool:
    v = (val or "").strip().lstrip("#").upper()
    if not re.fullmatch(r"[0-9A-F]{6}", v):
        return False
    r, g, b = int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16)
    return (r >= 200 and g <= 80 and b <= 80)


def is_red_run(run) -> bool:
    color = run.font.color
    if not color:
        return False
    rgb = color.rgb
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return (r >= 200 and g <= 80 and b <= 80)


def paragraph_text_and_is_red(paragraph: Paragraph) -> tuple[str, bool]:
    text = "".join(run.text for run in paragraph.runs).strip()
    any_red = any(is_red_run(run) and run.text.strip() for run in paragraph.runs)
    return text, any_red


def local(tag: str) -> str:
    return (tag or "").rsplit("}", 1)[-1]


def txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = []
        for t_node in r.iter():
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None):
                run_texts.append(t_node.text)
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue

        rpr = None
        for child in list(r):
            if str(getattr(child, "tag", "")).endswith("}rPr"):
                rpr = child
                break
        if rpr is None:
            continue

        color = None
        for c in rpr.iter():
            if str(getattr(c, "tag", "")).endswith("}color"):
                color = c
                break
        if color is None:
            continue
        val = None
        for k, v in getattr(color, "attrib", {}).items():
            if str(k).endswith("}val") or str(k) == "val":
                val = v
                break
        if val and is_red_hex(val):
            any_red = True

    return clean_text("".join(parts)), any_red


def textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    out: list[tuple[str, bool]] = []
    txbx_paras = []
    for el in p_elm.iter():
        if str(getattr(el, "tag", "")).endswith("}txbxContent"):
            for p2 in el.iter():
                if str(getattr(p2, "tag", "")).endswith("}p"):
                    txbx_paras.append(p2)
    for tx_p in txbx_paras:
        t, red = txbx_paragraph_text_and_is_red(tx_p)
        if t:
            out.append((t, red))
    return out


def drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    lines: list[tuple[str, bool]] = []
    seen = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue

        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = None
                for k, v in getattr(node, "attrib", {}).items():
                    if str(k).endswith("}val") or str(k) == "val":
                        val = v
                        break
                if val and is_red_hex(val):
                    any_red = True
        text = clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines


def iter_block_items(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def extract_items_with_red(docx_path: str, include_tables: bool = True) -> list[dict]:
    """
    Extract a stream of {text,is_red} lines from doc paragraphs + table cells.
    Also extracts text from textboxes/drawing shapes when anchored in paragraphs.
    """
    def extract_with_python_docx(path: str) -> list[dict]:
        doc = Document(path)
        items: list[dict] = []
        last = None

        def push(t: str, red: bool, src: str):
            nonlocal last
            t = clean_text(t)
            if not t:
                return
            key = (t, bool(red), src)
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red), "src": src})
            last = key

        def push_paragraph(p: Paragraph, src: str):
            t, red = paragraph_text_and_is_red(p)
            push(t, red, src)
            for t2, red2 in textbox_texts_in_paragraph(p):
                push(t2, red2, src)
            for t2, red2 in drawingml_texts_in_paragraph(p):
                push(t2, red2, src)

        def push_table(tbl: Table):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        push_paragraph(p, "table")
                    for sub in getattr(cell, "tables", []) or []:
                        push_table(sub)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                push_paragraph(block, "body")
            else:
                if include_tables:
                    push_table(block)
        return items

    def extract_with_xml_fallback(path: str) -> list[dict]:
        try:
            with zipfile.ZipFile(path) as z:
                raw = z.read("word/document.xml")
        except Exception:
            return []
        try:
            root = ET.fromstring(raw)
        except Exception:
            return []

        def run_is_red(run_el) -> bool:
            for el in run_el.iter():
                lname = local(getattr(el, "tag", ""))
                if lname == "color":
                    for k, v in getattr(el, "attrib", {}).items():
                        if str(k).endswith("}val") or str(k) == "val":
                            if v and is_red_hex(v):
                                return True
                if lname == "srgbClr":
                    for k, v in getattr(el, "attrib", {}).items():
                        if str(k).endswith("}val") or str(k) == "val":
                            if v and is_red_hex(v):
                                return True
            return False

        items: list[dict] = []
        last = None

        def push(parts: list[str], red: bool):
            nonlocal last
            t = clean_text("".join(parts))
            if not t:
                return
            key = (t, bool(red))
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red)})
            last = key

        def walk(node, in_p: bool = False, buf: list[str] | None = None, red_any: list[bool] | None = None):
            lname = local(str(getattr(node, "tag", "")))
            if lname == "p":
                b: list[str] = []
                r = [False]
                for ch in list(node):
                    walk(ch, in_p=True, buf=b, red_any=r)
                push(b, r[0])
                return
            if lname == "r":
                is_red = run_is_red(node)
                for ch in list(node):
                    walk(ch, in_p=in_p, buf=buf, red_any=red_any)
                if is_red and red_any is not None:
                    red_any[0] = True
                return
            if lname == "t":
                if in_p and buf is not None and getattr(node, "text", None):
                    buf.append(node.text)
                return
            for ch in list(node):
                walk(ch, in_p=in_p, buf=buf, red_any=red_any)

        for ch in list(root):
            walk(ch)
        return items

    def score(items: list[dict]) -> int:
        if not items:
            return 0
        qverb = re.compile(
            r"^(q\s*\d+\s*[\.)]\s*)?(list|describe|explain|outline|state|name|provide|define|identify|select|choose|pick|match|complete)\b",
            re.IGNORECASE,
        )
        sc = 0
        for it in items:
            t = clean_text(it.get("text", ""))
            if not t:
                continue
            if qverb.match(t) or "Which of the following" in t:
                sc += 3
            if t.endswith("?"):
                sc += 1
        return sc + min(80, len(items) // 8)

    # Prefer python-docx extraction so we can keep `src` ("body" vs "table") for downstream filtering.
    # XML fallback can't preserve this provenance reliably and can cause table content to leak into AI segmentation.
    items_docx = extract_with_python_docx(docx_path)
    return items_docx


HARD_QNUM_RE = re.compile(r"(?:(?<=\s)|^)(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
HARD_QNUM_RANGE_RE = re.compile(r"(?:(?<=\s)|^)(?:lo\s*)?q\s*\d+\s*[-–]\s*\d+\s*[\.\)]\s*", re.IGNORECASE)


def split_items_on_internal_qnums(items: list[dict]) -> list[dict]:
    out: list[dict] = []
    for it in items:
        t = clean_text(it.get("text", ""))
        if not t:
            continue
        src = it.get("src")
        if HARD_QNUM_RANGE_RE.search(t):
            out.append({"text": t, "is_red": bool(it.get("is_red")), "src": src})
            continue
        starts = [m.start() for m in HARD_QNUM_RE.finditer(t)]
        if len(starts) <= 1:
            out.append({"text": t, "is_red": bool(it.get("is_red")), "src": src})
            continue
        for a, b in zip(starts, starts[1:] + [len(t)]):
            seg = clean_text(t[a:b])
            if seg:
                out.append({"text": seg, "is_red": bool(it.get("is_red")), "src": src})
    return out


# ===================================================
# Table parsers (turn common assessor-guide tables into short-answer questions)
# ===================================================
def join_lines(lines: list[str]) -> str:
    parts = [clean_text(x) for x in (lines or []) if clean_text(x)]
    return "; ".join(parts).strip()


def table_to_grid(tbl: Table) -> list[list[list[str]]]:
    grid: list[list[list[str]]] = []
    for row in tbl.rows:
        r: list[list[str]] = []
        for cell in row.cells:
            lines = []
            for p in cell.paragraphs:
                t = clean_text(p.text)
                if t:
                    lines.append(t)
            r.append(lines)
        grid.append(r)
    return grid


def header_contains(grid: list[list[list[str]]], *needles: str) -> bool:
    if not grid or not grid[0]:
        return False
    header = " | ".join(join_lines(c) for c in grid[0]).lower()
    return all(n.lower() in header for n in needles)


def iter_tables_recursive(tbl: Table):
    yield tbl
    for row in tbl.rows:
        for cell in row.cells:
            for sub in getattr(cell, "tables", []) or []:
                yield from iter_tables_recursive(sub)


def iter_all_tables(doc: Document):
    for el in iter_block_items(doc):
        if isinstance(el, Table):
            yield from iter_tables_recursive(el)


def find_item_index(items: list[dict], needle: str) -> int | None:
    n = normalize_key(needle)
    if not n:
        return None
    for i, it in enumerate(items):
        t = normalize_key(it.get("text", ""))
        if not t:
            continue
        if t == n or n in t:
            return i
    return None


def parse_table_defined_terms_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    """
    Converts common 3-column tables like:
      Poultry ingredient | Definition | Style/method of cooking
    into per-row essay questions instead of Matching/MCQ.

    User preference:
    - These should become ONE short-answer question per row:
      "Define: <term>. Provide one style/method of cooking."
    """
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in iter_all_tables(doc):
        grid = table_to_grid(el)
        if not header_contains(grid, "poultry ingredient", "definition", "style"):
            continue

        for r in grid[1:]:
            if not r:
                continue
            term = join_lines(r[0] if len(r) > 0 else [])
            term = strip_q_prefix(term)
            term = clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term:
                continue
            if term.lower() in {"poultry ingredient", "definition", "style/method of cooking"}:
                continue
            k = normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)

            order = find_item_index(items, term) or find_item_index(items, f"Define {term}") or 10**9
            q = f"Define: {term}. Provide one style/method of cooking."
            out.append({"question": q, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": order, "qnum": None})
    return out


def parse_table_characteristics_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    """
    Converts tables like:
      Poultry type or cut | Essential characteristics
    into per-row essay questions.
    """
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in iter_all_tables(doc):
        grid = table_to_grid(el)
        if not header_contains(grid, "poultry type", "essential"):
            continue

        for r in grid[1:]:
            if not r:
                continue
            term = join_lines(r[0] if len(r) > 0 else [])
            term = strip_q_prefix(term)
            term = clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term:
                continue
            if term.lower() in {"poultry type or cut", "essential characteristics"}:
                continue
            k = normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)
            q = f"Describe the essential characteristics of: {term}."
            order = find_item_index(items, term) or find_item_index(items, f"Describe {term}") or 10**9
            out.append({"question": q, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": order, "qnum": None})
    return out


def collect_ignore_texts_from_forced_tables(docx_path: str) -> set[str]:
    """
    Some assessor guides use tables to show sample answers (often in red) like:
      Poultry ingredient | Definition | Style/method of cooking
      Poultry type or cut | Essential characteristics
      Classical chicken dishes | Contemporary chicken dishes

    We don't want the AI segmenter to turn those table cell values into extra questions/MCQs.
    """
    doc = Document(docx_path)
    ignore: set[str] = set()

    def add_lines(lines: list[str]):
        for ln in lines or []:
            t = clean_text(ln)
            if not t:
                continue
            # Avoid nuking very short single tokens like "Chicken" or "14"
            if len(t) < 12 and " " not in t:
                continue
            if re.fullmatch(r"\d+", t):
                continue
            ignore.add(t)

    for tbl in iter_all_tables(doc):
        grid = table_to_grid(tbl)
        if header_contains(grid, "poultry ingredient", "definition", "style"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
            continue
        if header_contains(grid, "poultry type", "essential"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
            continue
        if header_contains(grid, "classical chicken dishes", "contemporary chicken dishes"):
            for row in grid[1:]:
                for cell in row:
                    add_lines(cell)
            continue

    return ignore


# ===================================================
# Matching (rule-based; reliable for simple tables)
# ===================================================
MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)


def looks_like_matching_stem(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(MATCHING_STEM_RE.search(t2))


def table_fingerprint(grid: list[list[list[str]]]) -> str:
    rows = []
    for r in grid:
        rows.append("|".join(join_lines(c) for c in r))
    return normalize_key("||".join(rows))


def score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        left = join_lines(r[a])
        right = join_lines(r[b])
        if left and right:
            sc += 1
    return sc


def pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    if best_sc < 2:
        return None
    return best


def extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1):
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = join_lines(r[left_col])
        right = join_lines(r[right_col])
        if not left or not right:
            continue
        pairs.append({"left": left, "right": right})
    return pairs


def is_table_forced_essay(grid: list[list[list[str]]]) -> bool:
    return (
        header_contains(grid, "poultry ingredient", "definition", "style")
        or header_contains(grid, "poultry type", "essential")
        or header_contains(grid, "classical chicken dishes", "contemporary chicken dishes")
    )


def parse_matching_questions_doc_order(docx_path: str, items: list[dict] | None = None) -> list[dict]:
    doc = Document(docx_path)
    out = []
    seen = set()
    seq = 0

    def is_instructions_matching(pairs: list[dict], stem: str) -> bool:
        """
        Skip admin/instructions rubric tables that sometimes look like matching questions, e.g.
        'Match each Instructions to the correct For students...' with rows like:
          Range and conditions → Online; students must...
        These are not quiz questions.
        """
        s = normalize_key(stem or "")
        if "instructions" in s and ("for students" in s or "for learners" in s or "for assessors" in s):
            return True

        left_keys = {normalize_key((p.get("left") or "")) for p in (pairs or [])}
        common = {
            "range and conditions",
            "decision-making rules",
            "decision making rules",
            "pre-approved reasonable adjustments",
            "pre approved reasonable adjustments",
            "rubric",
            "instructions",
        }
        hit = sum(1 for k in common if k in left_keys)
        if hit >= 3:
            return True

        right_blob = normalize_key("; ".join(p.get("right") or "" for p in (pairs or [])))
        if "students must work through this assessment independently" in right_blob:
            return True
        if "false declarations may lead to withdrawal" in right_blob:
            return True
        if "feedback comments must be provided" in right_blob:
            return True
        return False

    for el in iter_all_tables(doc):
        seq += 1
        grid = table_to_grid(el)
        if is_table_forced_essay(grid):
            continue

        fp = table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)

        cols = pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue

        header = grid[0] if grid else []
        hL = (join_lines(header[left_col]) if header and left_col < len(header) else "Left")
        hR = (join_lines(header[right_col]) if header and right_col < len(header) else "Right")
        stem = f"Match each '{hL}' to the correct '{hR}'."

        if is_instructions_matching(pairs, stem):
            continue

        order = seq
        if items:
            for cand in (hL, hR, pairs[0].get("left", ""), pairs[0].get("right", "")):
                idx = find_item_index(items, cand)
                if idx is not None:
                    order = idx
                    break

        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False,
            "_order": order,
            "qnum": None,
        })

    return out


# ===================================================
# OpenAI (AI segmentation for MCQ/Essay)
# ===================================================
@dataclass
class OpenAIConfig:
    api_key: str
    model: str
    base_url: str = "https://api.openai.com"
    timeout_s: int = 120


def openai_responses_json_schema(prompt: str, schema_name: str, schema: dict, cfg: OpenAIConfig) -> tuple[dict | None, str | None]:
    url = cfg.base_url.rstrip("/") + "/v1/responses"
    headers = {
        "Authorization": f"Bearer {cfg.api_key}",
        "Content-Type": "application/json",
    }
    body = {
        "model": cfg.model,
        "input": prompt,
        "text": {
            "format": {
                "type": "json_schema",
                "name": schema_name,
                "schema": schema,
                "strict": True,
            }
        },
    }
    try:
        r = requests.post(url, headers=headers, json=body, timeout=cfg.timeout_s)
    except Exception as e:
        return None, f"OpenAI request failed: {e}"
    if r.status_code >= 400:
        return None, f"OpenAI error {r.status_code}: {r.text}"
    try:
        data = r.json()
    except Exception as e:
        return None, f"OpenAI JSON parse failed: {e}"
    try:
        out = data["output"][0]["content"][0]
        if out.get("type") == "output_text" and out.get("text"):
            return json.loads(out["text"]), None
        if out.get("type") == "output_json" and out.get("json"):
            return out["json"], None
        if "text" in out:
            return json.loads(out["text"]), None
    except Exception as e:
        return None, f"OpenAI response parse failed: {e}"
    return None, "OpenAI returned an unexpected response shape."


IGNORE_LINE_RE = re.compile(
    r"^(for learners|for assessors|for students|range and conditions|decision-making rules|pre-approved|rubric|feedback|knowledge test)\b",
    re.IGNORECASE,
)
IGNORE_SECTION_RE = re.compile(
    r"^\s*(?:q\s*\d+\s*[-–]\s*\d+[\.\)]\s*)?(?:the\s+following\s+questions\s+require\s+you\s+to|the\s+following\s+questions\s+require)\b",
    re.IGNORECASE,
)
IGNORE_TABLE_RE = re.compile(
    r"^(poultry ingredient|definition|style/method of cooking|poultry type or cut|essential characteristics|classical chicken dishes|contemporary chicken dishes)\b",
    re.IGNORECASE,
)

COOKERY_METHOD_WORD_RE = re.compile(
    r"^(?:pan[-\s]?fry|deep[-\s]?fry|stir[-\s]?fry|roast|bake|grill|bbq|braise|stew|simmer|poach|saute|sauté|steam|boil)\b",
    re.IGNORECASE,
)

QUESTION_START_RE = re.compile(
    r"^(?:"
    r"(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*"
    r")?"
    r"(?:critically\s+)?"
    r"(?:"
    r"which of the following|"
    r"select|choose|pick|match|complete|"
    r"list|name|identify|define|describe|explain|outline|state|provide|"
    r"illustrate|evaluate|determine|articulate|discuss|analyse|analyze|compare|review|appraise|"
    r"assess|"
    r"what|when|where|why|how|"
    r"must\b"
    r")\b",
    re.IGNORECASE,
)


def looks_like_question_start(text: str) -> bool:
    t = clean_text(text)
    if not t:
        return False
    if t.endswith("?"):
        return True
    if "____" in t or "___" in t:
        return True
    return bool(QUESTION_START_RE.match(t))


def looks_like_answer_guide_bullet(text: str) -> bool:
    t = clean_text(text)
    if not t:
        return False
    if t.lower().startswith(("answer may address", "answer must address", "answer needs to address")):
        return True
    # Common rubric bullets that should never become questions/options
    if t.lower() in {"that is blank", "has nothing written in the space provided"}:
        return True
    if t.lower().startswith("does not attempt to answer"):
        return True
    return False


OPTION_LINE_RE = re.compile(
    r"^\s*(?:"
    r"(?:option\s*\d+)|"
    r"(?:\(?[a-h]\)|[a-h][\.\)])|"
    r"(?:\(?i{1,3}v?\)|i{1,3}v?[\.\)])"
    r")\s+",
    re.IGNORECASE,
)


def looks_like_option_line(text: str) -> bool:
    t = clean_text(text)
    if not t:
        return False
    if OPTION_LINE_RE.match(t):
        return True
    # Common MCQ option patterns that show up without letters in some docs:
    if t.lower().startswith(("true", "false")):
        return True
    return False


def is_admin_or_meta_line(text: str) -> bool:
    t = clean_text(text)
    if not t:
        return True
    tl = t.lower()
    # Common instruction lines that should never become questions
    if tl.startswith("when you have completed all questions"):
        return True
    if tl.startswith("by submitting your"):
        return True
    if tl.startswith("where a learner is assessed as"):
        return True
    if IGNORE_LINE_RE.match(t) or IGNORE_SECTION_RE.match(t) or IGNORE_TABLE_RE.match(t):
        return True
    if looks_like_answer_guide_bullet(t):
        return True
    if ANSWER_GUIDE_START_RE.match(t):
        return True
    return False


def parse_essay_questions_rule_based(items: list[dict]) -> list[dict]:
    """
    Fallback: when AI misses short-answer (essay) questions, extract them directly from the item stream.
    Conservative: only adds a question if it looks like a question stem and is not followed by option-like lines.
    Also tries hard to avoid turning "Answer needs to address..." bullet points into separate questions.
    """
    out: list[dict] = []
    seen: set[str] = set()
    in_answer_guide = False

    def has_answer_guide_soon(idx: int) -> bool:
        for j in range(idx + 1, min(len(items), idx + 8)):
            t2 = clean_text(items[j].get("text", ""))
            if not t2:
                continue
            if ANSWER_GUIDE_START_RE.match(t2) or ANSWER_GUIDE_ANY_RE.search(t2):
                return True
        return False

    i = 0
    while i < len(items):
        t = clean_text(items[i].get("text", ""))
        if not t or is_admin_or_meta_line(t):
            i += 1
            continue

        # In these assessor guides, red text is almost always rubric/sample-answer content.
        # Avoid turning any red line into a question stem.
        if bool(items[i].get("is_red")):
            i += 1
            continue

        # Track answer-guide blocks so we don't treat bullets as questions.
        if ANSWER_GUIDE_START_RE.match(t) or ANSWER_GUIDE_ANY_RE.search(t):
            in_answer_guide = True
            i += 1
            continue

        # If merged with "Answer must address", keep only stem.
        if ANSWER_GUIDE_ANY_RE.search(t):
            t = strip_answer_guide(t)
        stem = trim_after_question_mark(strip_q_prefix(t))
        stem = trim_after_sentence_if_long(stem)
        if not stem or len(stem) < 10 or not looks_like_question_start(stem):
            i += 1
            continue

        # If we're inside an answer-guide block, only exit when we see a "real" new question:
        # either ends with '?' or is followed soon by an Answer-guide marker.
        if in_answer_guide and not (stem.endswith("?") or has_answer_guide_soon(i)):
            i += 1
            continue

        # If the next few lines look like options, don't treat as essay.
        optionish = 0
        for j in range(i + 1, min(len(items), i + 8)):
            t2 = clean_text(items[j].get("text", ""))
            if not t2:
                continue
            if is_admin_or_meta_line(t2):
                continue
            if looks_like_option_line(t2):
                optionish += 1
        if optionish >= 2:
            i += 1
            continue

        # For non-'?' stems, require an answer-guide marker soon; this avoids adding rubric bullets.
        if not stem.endswith("?") and not has_answer_guide_soon(i):
            i += 1
            continue

        k = normalize_key(stem)
        if k and k not in seen:
            seen.add(k)
            out.append({
                "question": stem,
                "options": [],
                "correct": [],
                "multi": False,
                "kind": "essay",
                "_order": i,
                "qnum": None,
            })

        in_answer_guide = has_answer_guide_soon(i)
        i += 1
    return out


def filter_items_for_ai(
    items: list[dict],
    ignore_terms: set[str] | None = None,
    ignore_texts: set[str] | None = None,
    mode: str = "balanced",
) -> list[dict]:
    out: list[dict] = []
    ignore_terms = ignore_terms or set()
    ignore_terms_norm = {normalize_key(t) for t in ignore_terms if normalize_key(t)}
    ignore_term_prefixes = sorted([t for t in ignore_terms_norm if t], key=len, reverse=True)
    ignore_texts = ignore_texts or set()
    ignore_texts_norm = {normalize_key(t) for t in ignore_texts if normalize_key(t)}
    mode = (mode or "balanced").strip().lower()
    if mode not in {"balanced", "loose", "strict"}:
        mode = "balanced"
    in_answer_guide = False
    for it in items:
        t = clean_text(it.get("text", ""))
        if not t:
            continue
        if normalize_key(t) in ignore_texts_norm:
            continue
        if IGNORE_LINE_RE.match(t):
            continue
        if IGNORE_SECTION_RE.match(t):
            continue
        if mode in {"balanced", "strict"}:
            if ANSWER_GUIDE_START_RE.match(t):
                in_answer_guide = True
                continue
            if ANSWER_GUIDE_ANY_RE.search(t):
                # Sometimes a question and "Answer must address..." are merged into one line.
                # Keep the question part, but enter answer-guide mode for following lines.
                pre = strip_answer_guide(t)
                if pre and looks_like_question_start(pre) and len(pre) >= 10:
                    out.append({"text": pre, "is_red": False})
                in_answer_guide = True
                continue
        if IGNORE_TABLE_RE.match(t):
            continue
        if mode in {"balanced", "strict"} and looks_like_answer_guide_bullet(t):
            continue
        # Ignore table row labels (prevents AI turning table rows into MCQ/extra questions)
        tn = normalize_key(t)
        if tn in ignore_terms_norm:
            continue
        # Also skip lines that start with a known table term and then include other columns
        # e.g. "Giblets Chicken offal comprising heart, liver, kidneys"
        skipped = False
        for pref in ignore_term_prefixes:
            if tn.startswith(pref + " "):
                skipped = True
                break
        if skipped:
            continue
        # Handles "Q12. Fillets" style row labels
        m = re.match(r"^\s*(?:q\s*\d+\s*[\.\)]\s*)?(.*)$", t, flags=re.IGNORECASE)
        if m:
            rest = normalize_key(clean_text(m.group(1)))
            if rest in ignore_terms_norm:
                continue

        # While inside an answer-guide section, skip bullets/sample answers until next question starts.
        if mode == "strict" and in_answer_guide:
            if looks_like_question_start(t):
                in_answer_guide = False
            else:
                # Most answer-guide bullets are red and/or short phrases (including table content).
                if bool(it.get("is_red")):
                    continue
                if len(t) <= 120 and not t.endswith("?"):
                    continue
                continue
        if mode == "balanced" and in_answer_guide:
            if looks_like_question_start(t):
                in_answer_guide = False
            else:
                # Balanced: still skip most answer-guide bullets, but be less aggressive.
                if len(t) <= 80 and not t.endswith("?"):
                    continue
                if bool(it.get("is_red")) and len(t) <= 80:
                    continue
                continue

        # Ignore lone cookery-method bullets (common in answer-guide tables)
        if mode in {"balanced", "strict"} and len(t) <= 20 and COOKERY_METHOD_WORD_RE.match(t):
            continue

        out.append(it)
    return out


def ai_segment_items_openai(items: list[dict], cfg: OpenAIConfig) -> tuple[list[dict], list[str]]:
    log: list[str] = []
    if not items:
        return [], log

    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "kind": {"type": "string", "enum": ["mcq", "essay"]},
                        "stem": {"type": "array", "items": {"type": "integer"}},
                        "options": {"type": "array", "items": {"type": "array", "items": {"type": "integer"}}},
                    },
                    "required": ["kind", "stem", "options"],
                },
            }
        },
        "required": ["questions"],
    }

    def to_line(i: int) -> str:
        t = clean_text(items[i].get("text", ""))
        red = "R1" if items[i].get("is_red") else "R0"
        return f"I{i}|{red}|{t}"

    max_block_items = 170
    overlap = 50
    blocks = []
    start = 0
    while start < len(items):
        end = min(len(items), start + max_block_items)
        blocks.append((start, end))
        if end >= len(items):
            break
        start = max(0, end - overlap)

    all_qs: list[dict] = []

    def should_demote_mcq_to_essay(stem_text: str, options: list[str], correct: list[int]) -> bool:
        """
        In assessor guides, we sometimes see short-answer questions where the *answer* appears as a single
        red line (or the model collapses options). These should stay Essay/Short Answer, not MCQ.
        """
        s = normalize_key(stem_text)
        if not s:
            return False
        # If options collapsed to 1 (or 2 with one being a duplicate-ish), treat as short answer.
        if len(options) <= 1:
            return True
        if len(options) == 2:
            a, b = normalize_key(options[0]), normalize_key(options[1])
            if a and b and (a in b or b in a):
                return True
        # "What is/was the name..." + tiny answer list typically means short answer with sample answer.
        if s.startswith(("what is the name", "what was the name", "what is meant", "what is the origin")):
            if len(options) <= 3 and len(correct) <= 1:
                return True
        # If any option is basically repeating the stem, it's not a real MCQ.
        for opt in options:
            o = normalize_key(opt)
            if o and len(o) > 25 and (o in s or s in o):
                return True
        return False
    for (a, b) in blocks:
        ctx = []
        for i in range(a, b):
            t = clean_text(items[i].get("text", ""))
            if not t:
                continue
            ctx.append(to_line(i))
        if len(ctx) < 6:
            continue

        log.append(f"AI block: {a}-{b} lines={len(ctx)}")
        prompt = (
            "You are segmenting a DOCX extraction into Canvas quiz questions.\n"
            "Return STRICT JSON only (per schema).\n"
            "\n"
            "Hard rules:\n"
            "- You MUST NOT invent any text.\n"
            "- You may ONLY reference item indices (I<n>) from the provided list.\n"
            "- Keep original order (earlier indices first).\n"
            "- Do NOT create questions from instructions/policy/rubric.\n"
            "- Do NOT include assessor guide content like 'Answer may/must/needs address' or sample answers in stems/options.\n"
            "- Correct MCQ options are those with R1 (red). For essay questions: options must be [].\n"
            "\n"
            "MCQ rules:\n"
            "- Options are typically lettered (a), (b), etc or separate lines under a prompt.\n"
            "- If you cannot find at least 2 options, do NOT output an MCQ.\n"
            "\n"
            "Essay rules:\n"
            "- Use the question prompt only.\n"
            "\n"
            "Items (format: I<index>|R0/R1|text):\n"
            + "\n".join(ctx)
        )

        data, err = openai_responses_json_schema(prompt, "segment_questions", schema, cfg)
        if err:
            log.append(f"  block failed: {err}")
            continue
        qs = data.get("questions") if isinstance(data, dict) else None
        if not isinstance(qs, list):
            log.append("  block skipped: missing questions[]")
            continue

        for q in qs:
            if not isinstance(q, dict):
                continue
            kind = (q.get("kind") or "").strip().lower()
            stem_ids = q.get("stem") if isinstance(q.get("stem"), list) else []
            if kind not in ("mcq", "essay"):
                continue
            if not stem_ids or not all(isinstance(x, int) for x in stem_ids):
                continue
            if any(x < 0 or x >= len(items) for x in stem_ids):
                continue

            stem_text = clean_text(" ".join(clean_text(items[x].get("text", "")) for x in stem_ids))
            stem_text = strip_q_prefix(strip_answer_guide(stem_text))
            stem_text = trim_after_question_mark(stem_text)
            stem_text = trim_after_sentence_if_long(stem_text)
            if not stem_text or len(stem_text) < 10:
                continue
            if stem_text.lower().startswith(("answer may address", "answer must address", "answer needs to address")):
                continue
            if not looks_like_question_start(stem_text):
                # Prevent creating questions from answer-guide bullets like "Appearance and presentation"
                continue

            if kind == "essay":
                all_qs.append({
                    "question": stem_text,
                    "options": [],
                    "correct": [],
                    "multi": False,
                    "kind": "essay",
                    "_order": min(stem_ids),
                    "qnum": None,
                })
                continue

            opt_groups = q.get("options") if isinstance(q.get("options"), list) else []
            option_texts: list[str] = []
            correct: list[int] = []
            for group in opt_groups:
                if not isinstance(group, list) or not group or not all(isinstance(x, int) for x in group):
                    continue
                if any(x < 0 or x >= len(items) for x in group):
                    continue
                t = clean_text(" ".join(clean_text(items[x].get("text", "")) for x in group))
                if not t:
                    continue
                if ANSWER_GUIDE_START_RE.match(t) or IGNORE_TABLE_RE.match(t) or IGNORE_LINE_RE.match(t):
                    continue
                option_texts.append(t)
                if any(bool(items[x].get("is_red")) for x in group):
                    correct.append(len(option_texts) - 1)

            seen = set()
            out_opts = []
            out_corr = []
            for i_opt, opt in enumerate(option_texts):
                k = normalize_key(opt)
                if k in seen:
                    continue
                seen.add(k)
                if i_opt in correct:
                    out_corr.append(len(out_opts))
                out_opts.append(opt)
            if len(out_opts) < 2:
                continue

            if should_demote_mcq_to_essay(stem_text, out_opts, out_corr):
                all_qs.append({
                    "question": stem_text,
                    "options": [],
                    "correct": [],
                    "multi": False,
                    "kind": "essay",
                    "_order": min(stem_ids),
                    "qnum": None,
                })
            else:
                all_qs.append({
                    "question": stem_text,
                    "options": out_opts,
                    "correct": out_corr,
                    "multi": ("apply" in stem_text.lower()) or (len(out_corr) > 1),
                    "kind": "mcq",
                    "_order": min(stem_ids),
                    "qnum": None,
                })

    # de-dupe by question text (keep earliest order)
    deduped = []
    seen_q = set()
    for q in sorted(all_qs, key=lambda q: int(q.get("_order", 10**9))):
        k = normalize_key(q.get("question", ""))
        if not k or k in seen_q:
            continue
        seen_q.add(k)
        deduped.append(q)
    return deduped, log


# ===================================================
# Post-processing helpers
# ===================================================
def question_dedupe_key(q: dict) -> str:
    kind = (q.get("kind") or "").lower().strip()
    if kind == "matching":
        pairs = q.get("pairs") or []
        parts = []
        for p in pairs:
            left = normalize_key((p or {}).get("left") or "")
            right = normalize_key((p or {}).get("right") or "")
            if left or right:
                parts.append(f"{left}->{right}")
        return "matching|" + normalize_key(q.get("question", "")) + "|" + "|".join(parts)
    if kind == "mcq":
        opts = [normalize_key(o) for o in (q.get("options") or []) if normalize_key(o)]
        return "mcq|" + normalize_key(q.get("question", "")) + "|" + "|".join(opts)
    # essay/short answer
    return "essay|" + normalize_key(q.get("question", ""))


def dedupe_questions(questions: list[dict]) -> tuple[list[dict], int]:
    out: list[dict] = []
    seen: set[str] = set()
    removed = 0
    for q in questions:
        k = question_dedupe_key(q)
        if not k or k in seen:
            removed += 1
            continue
        seen.add(k)
        out.append(q)
    return out, removed


# ===================================================
# UI
# ===================================================
st.set_page_config(page_title="Canvas Quiz Builder (AI-only v3)", layout="wide")
st.title("Canvas Quiz Builder (AI-only v3)")
st.caption(f"Build: {BUILD_ID}  |  Runs on port 8505 via /Users/jargalmaa/Downloads/web_ui/run_canvas_ai_v3.sh")

st.sidebar.header("AI Settings")
openai_key = st.sidebar.text_input("OpenAI API key", value=os.getenv("OPENAI_API_KEY", ""), type="password")
openai_model = st.sidebar.text_input("Model", value=os.getenv("OPENAI_MODEL", "gpt-4.1-mini"))
openai_base_url = st.sidebar.text_input("Base URL", value=os.getenv("OPENAI_BASE_URL", "https://api.openai.com"))
st.sidebar.divider()
ignore_tables = st.sidebar.checkbox("Ignore tables (skip table text/questions)", value=False)
include_table_essays = st.sidebar.checkbox("Convert common tables to short-answer questions", value=True)
include_rule_essay_fallback = st.sidebar.checkbox("Add rule-based short-answer fallback", value=True)
ai_filter_mode = st.sidebar.selectbox(
    "AI filter mode",
    ["balanced", "loose", "strict"],
    index=0,
    help="Loose keeps more lines (may add noise). Strict removes more answer-guide text (can miss questions if the DOCX is messy).",
)
st.sidebar.caption("AI segments MCQ/Essay. Matching is still detected from tables (rule-based).")

docx_path_input = st.text_input(
    "DOCX path (optional)",
    value="",
    placeholder="/Users/jargalmaa/Downloads/quizzes/SITHCCC035 Knowledge Test Assessor Guide (Short Answer).docx",
)
uploaded = st.file_uploader("...or upload DOCX", type=["docx"])
parse_btn = st.button("Parse DOCX", type="primary")
log_box = st.empty()

if "questions" not in st.session_state:
    st.session_state.questions = []

if parse_btn:
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        docx_path = ""
        if docx_path_input.strip():
            docx_path = docx_path_input.strip()
            if not os.path.exists(docx_path):
                raise RuntimeError(f"DOCX path not found: {docx_path}")
        elif uploaded:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(uploaded.read())
                docx_path = tmp.name
        else:
            raise RuntimeError("Provide a DOCX path or upload a DOCX file.")

        items = extract_items_with_red(docx_path, include_tables=(not ignore_tables))
        items = split_items_on_internal_qnums(items)

        matching = [] if ignore_tables else parse_matching_questions_doc_order(docx_path, items)

        table_essays: list[dict] = []
        if include_table_essays and not ignore_tables:
            table_essays.extend(parse_table_defined_terms_as_essays(docx_path, items))
            table_essays.extend(parse_table_characteristics_as_essays(docx_path, items))

        # Build a set of terms we already converted from tables so AI doesn't create extra questions/MCQs from them.
        ignore_terms: set[str] = set()
        for q in table_essays:
            qt = clean_text(q.get("question", ""))
            m = re.match(r"^Define:\s*(.+?)\s*\.", qt, flags=re.IGNORECASE)
            if m:
                ignore_terms.add(clean_text(m.group(1)))
                continue
            m = re.match(r"^Describe the essential characteristics of:\s*(.+?)\s*\.", qt, flags=re.IGNORECASE)
            if m:
                ignore_terms.add(clean_text(m.group(1)))

        ignore_texts: set[str] = set()
        if not ignore_tables:
            ignore_texts = collect_ignore_texts_from_forced_tables(docx_path)

        ai_log = []
        ai_qs: list[dict] = []
        if not openai_key.strip():
            ai_log.append("OpenAI API key missing.")
        else:
            cfg = OpenAIConfig(
                api_key=openai_key.strip(),
                model=openai_model.strip(),
                base_url=openai_base_url.strip() or "https://api.openai.com",
            )
            ai_input = filter_items_for_ai(items, ignore_terms=ignore_terms, ignore_texts=ignore_texts, mode=ai_filter_mode)
            ai_qs, ai_log = ai_segment_items_openai(ai_input, cfg)

        rule_essays: list[dict] = []
        if include_rule_essay_fallback:
            rule_essays = parse_essay_questions_rule_based(items)

        questions = []
        questions.extend(matching)
        questions.extend(table_essays)
        questions.extend(ai_qs)
        questions.extend(rule_essays)
        questions.sort(key=lambda q: int(q.get("_order", 10**9)))
        questions, removed_dupes = dedupe_questions(questions)

        st.session_state.questions = questions

        print("DEBUG: build:", BUILD_ID)
        print("DEBUG: ignore_tables:", bool(ignore_tables))
        print("DEBUG: include_table_essays:", bool(include_table_essays))
        print("DEBUG: include_rule_essay_fallback:", bool(include_rule_essay_fallback))
        print("DEBUG: ai_filter_mode:", ai_filter_mode)
        print(
            "DEBUG: items extracted:",
            len(items),
            "ai_input:",
            len(filter_items_for_ai(items, ignore_terms=ignore_terms, ignore_texts=ignore_texts, mode=ai_filter_mode)),
        )
        print(
            "DEBUG: matching:",
            len(matching),
            "table_essays:",
            len(table_essays),
            "ai_mcq:",
            sum(1 for q in ai_qs if q.get("kind") == "mcq"),
            "ai_essay:",
            sum(1 for q in ai_qs if q.get("kind") == "essay"),
        )
        print("DEBUG: rule_essays:", len(rule_essays))
        print("DEBUG: removed_dupes:", removed_dupes)
        print("Parsed questions:", len(questions))
        for line in ai_log[:80]:
            print(line)

    log_box.code(buf.getvalue())
    st.success(f"✅ Parsed {len(st.session_state.questions)} questions.")

questions = st.session_state.questions or []
if questions:
    st.subheader("Questions")
    colp1, colp2 = st.columns([1, 1])
    page_size = colp1.selectbox("Questions per page", [5, 10, 15, 20, 30], index=1)
    total = len(questions)
    total_pages = max(1, math.ceil(total / page_size))
    page = colp2.number_input("Page", min_value=1, max_value=total_pages, value=1, step=1)
    start = (page - 1) * page_size
    end = min(start + page_size, total)
    st.caption(f"Showing questions {start+1}–{end} of {total}")

    for i in range(start, end):
        q = questions[i]
        kind = (q.get("kind") or "").lower()
        label_kind = "Matching" if kind == "matching" else ("Essay/Short Answer" if kind == "essay" else "MCQ")
        preview = strip_q_prefix(q.get("question", ""))[:90]
        with st.expander(f"Q{i+1} ({label_kind}): {preview}"):
            st.write(q.get("question", ""))
            if kind == "mcq":
                st.write("Options:")
                for j, opt in enumerate(q.get("options") or []):
                    mark = "✅" if j in (q.get("correct") or []) else ""
                    st.write(f"- {opt} {mark}")
            if kind == "matching":
                st.write("Pairs:")
                for p in q.get("pairs") or []:
                    st.write(f"- {p.get('left','')} → {p.get('right','')}")
