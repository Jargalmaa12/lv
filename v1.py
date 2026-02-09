# ==========================================================
# Learning Vault – Canvas Quiz Uploader (FULL STABLE VERSION)
# ✅ Matching + MCQ + Essay
# ✅ Matching fixed: no duplicates, correct stems, list-cells, no “wrong table”
# ✅ No undefined functions / no duplicate parsers
# ✅ Clears previous file results properly (no “previous answers” carry over)
# ==========================================================

import os
import re
import io
import math
import hashlib
import tempfile
import contextlib
from datetime import datetime, date, time

import pytz
import requests
import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

# ===================================================
# TIMEZONE
# ===================================================
TZ_NAME = "Australia/Sydney"
tz = pytz.timezone(TZ_NAME)

def combine_date_time(d: date | None, t: time | None):
    if not d or not t:
        return None
    dt = tz.localize(datetime.combine(d, t))
    return dt.isoformat()

# ===================================================
# UI
# ===================================================
st.set_page_config(page_title="Learning Vault – Canvas Quiz Uploader", layout="wide")
st.title("Learning Vault – Canvas Quiz Uploader (Web Interface)")

# ===================================================
# SESSION STATE DEFAULTS
# ===================================================
def ss_init(key, value):
    if key not in st.session_state:
        st.session_state[key] = value

ss_init("logged_in", False)
ss_init("me", None)
ss_init("canvas_token", "")
ss_init("canvas_base_url", "https://learningvault.test.instructure.com/api/v1")
ss_init("courses_cache", None)
ss_init("selected_course_id", None)

ss_init("docx_filename", None)
ss_init("description_html", "")
ss_init("questions", [])
ss_init("parsed_ok", False)
ss_init("parse_run_id", 0)


ss_init("details", {
    "shuffle_answers": True,
    "time_limit": 0,
    "allow_multiple_attempts": False,
    "allowed_attempts": 2,
    "scoring_policy": "keep_highest",
    "one_question_at_a_time": False,
    "show_correct_answers": False,
    "access_code_enabled": False,
    "access_code": "",
    "due_at": "",
    "unlock_at": "",
    "lock_at": "",
})

# ===================================================
# HELPERS
# ===================================================
def clean_text(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "")).strip()

def normalize_key(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s

Q_PREFIX_RE = re.compile(r"^[^A-Za-z0-9]*(?:lo\s*)?Q\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
def strip_q_prefix(line: str) -> str:
    return Q_PREFIX_RE.sub("", (line or "").strip()).strip()

def question_fingerprint(q: dict) -> str:
    qt = normalize_key(q.get("question", ""))
    kind = normalize_key(q.get("kind", ""))
    opts = [normalize_key(x) for x in (q.get("options") or [])]
    pairs = q.get("pairs") or []
    pairs_blob = "||".join([normalize_key(p.get("left","")) + "=>" + normalize_key(p.get("right","")) for p in pairs])
    blob = kind + "||" + qt + "||" + "||".join(opts) + "||" + pairs_blob
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()

def dedupe_questions(questions: list) -> list:
    seen = set()
    out = []
    for q in questions:
        fp = question_fingerprint(q)
        if fp in seen:
            continue
        seen.add(fp)
        out.append(q)
    return out

# ===================================================
# RED DETECTION (correct answers in red)
# ===================================================
def is_red_run(run) -> bool:
    color = run.font.color
    if not color:
        return False
    rgb = color.rgb
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return (r >= 200 and g <= 80 and b <= 80)

def paragraph_text_and_is_red(paragraph):
    text = "".join(run.text for run in paragraph.runs).strip()
    any_red = any(is_red_run(run) and run.text.strip() for run in paragraph.runs)
    return text, any_red

# ===================================================
# DOCX ORDER ITERATOR (TOP LEVEL ONLY)
# Important: no nested recursion here → stops “2 matching became 6”
# ===================================================
def iter_block_items(doc):
    """
    Yield Paragraph and Table objects in true top-level document order.
    SAFE: does not rely on isinstance(Document).
    """
    body = doc.element.body

    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


# ===================================================
# EXTRACT ITEMS (paragraph + table cells in true top-level order)
# for MCQ/Essay parsing + description block
# ===================================================
def extract_items_with_red(docx_path):
    doc = Document(docx_path)
    items = []

    def push_text(t: str, is_red: bool):
        t = clean_text(t)
        if t:
            items.append({"text": t, "is_red": is_red})

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t, red = paragraph_text_and_is_red(block)
            push_text(t, red)
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t, red = paragraph_text_and_is_red(p)
                        push_text(t, red)

    return items

# ===================================================
# QUESTION / NOISE RULES
# ===================================================
NOISE_RE = re.compile(
    r"^(Instructions|For learners|For assessors|Range and conditions|Decision-making rules|"
    r"Pre-approved reasonable adjustments|Rubric|Knowledge Test)\b",
    re.IGNORECASE
)

QUESTION_CMD_INNER_RE = re.compile(
    r"\b(Which\s+of\s+the\s+following\b|"
    r"(Identify|Select)\s+(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b)",
    re.IGNORECASE
)

COMMAND_QUESTION_RE = re.compile(
    r"^(Illustrate|Critically\s+(?:assess|analyse|analyze|evaluate)|"
    r"Evaluate|Determine|Articulate|Prescribe|Analyse|Analyze|Review|Recommend)\b.+",
    re.IGNORECASE
)

RUBRIC_START_RE = re.compile(r"^Answer\s+needs\s+to\s+address\b", re.IGNORECASE)
ESSAY_GUIDE_RE = re.compile(r"^Answer\s+(may|must)\s+address", re.IGNORECASE)

# ===================================================
# BUILD DESCRIPTION HTML (For learners block)
# ===================================================
def build_description(items):
    collecting = False
    lines = []

    for it in items:
        t = clean_text(it.get("text",""))
        if not t:
            continue

        if re.search(r"\bFor learners\b", t, re.IGNORECASE):
            collecting = True

        if collecting and (
            QUESTION_CMD_INNER_RE.search(t)
            or COMMAND_QUESTION_RE.match(strip_q_prefix(t))
            or re.search(r"\bdragging\s+and\s+dropping\b|\bdrag\s+and\s+drop\b|\bComplete\s+the\s+table\b", t, re.IGNORECASE)
        ):
            break

        if collecting:
            lines.append(t)

    if not lines:
        return ""

    html_parts = []
    in_list = False

    for ln in lines:
        ln = ln.strip()

        if "•" in ln and not ln.strip().startswith("•"):
            before, *bullets = [p.strip() for p in ln.split("•") if p.strip()]
            if before:
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                html_parts.append(f"<p>{before}</p>")

            if bullets:
                if not in_list:
                    html_parts.append("<ul>")
                    in_list = True
                for b in bullets:
                    html_parts.append(f"<li>{b}</li>")
            continue

        if ln.startswith("•"):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{ln.lstrip('•').strip()}</li>")
            continue

        if in_list:
            html_parts.append("</ul>")
            in_list = False

        html_parts.append(f"<p>{ln}</p>")

    if in_list:
        html_parts.append("</ul>")

    return "\n".join(html_parts)

# ===================================================
# PARSE MCQ QUESTIONS (options + correct in red)
# ===================================================
def parse_mcq_questions(items):
    questions_list = []
    current_q = None
    current_opts = []

    def flush():
        nonlocal current_q, current_opts
        if not current_q:
            return

        opts = [o for o in current_opts if not NOISE_RE.match(o["text"])]
        option_texts = [o["text"] for o in opts]
        correct = [i for i, o in enumerate(opts) if o["is_red"]]

        qtext = strip_q_prefix(current_q.strip())
        qlower = qtext.lower()
        multi = bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower)) or ("apply" in qlower) or (len(correct) > 1)

        questions_list.append({
            "question": qtext,
            "options": option_texts,
            "correct": correct,
            "multi": multi,
            "kind": "mcq"
        })

        current_q = None
        current_opts = []

    for it in items:
        line = clean_text(it.get("text",""))
        if not line:
            continue
        if NOISE_RE.match(line):
            continue
        if ESSAY_GUIDE_RE.match(line):
            current_q = None
            current_opts = []
            continue

        m = QUESTION_CMD_INNER_RE.search(line)
        if m:
            flush()
            start = m.start()
            stem = line[:start].strip()
            cmd_plus = line[start:].strip()
            q_line = f"{stem} {cmd_plus}".strip() if stem else cmd_plus
            current_q = strip_q_prefix(q_line)
            current_opts = []
            continue

        if current_q:
            current_opts.append({"text": line, "is_red": it.get("is_red", False)})

    flush()
    return [q for q in questions_list if len(q.get("options") or []) >= 2 and len(q.get("question") or "") >= 10]

# ===================================================
# PARSE ESSAY QUESTIONS
# ===================================================
def parse_essay_questions(items):
    questions = []
    n = len(items)

    i = 0
    while i < n:
        raw = clean_text(items[i].get("text",""))
        if not raw or NOISE_RE.match(raw):
            i += 1
            continue

        line = strip_q_prefix(raw)

        if COMMAND_QUESTION_RE.match(line):
            j = i + 1
            next_line = ""
            while j < n:
                nxt = clean_text(items[j].get("text",""))
                if nxt and not NOISE_RE.match(nxt):
                    next_line = nxt
                    break
                j += 1

            if RUBRIC_START_RE.match(next_line):
                questions.append({
                    "question": line,
                    "options": [],
                    "correct": [],
                    "multi": False,
                    "kind": "essay"
                })
                i = j + 1
                continue

        i += 1

    return [q for q in questions if len((q.get("question") or "").strip()) >= 10]


# ===================================================
# MATCHING PARSER (FULL)  ✅ paste this BEFORE you call it
# ===================================================

# ===================================================
# MATCHING PARSER (RECUSRIVE + DEDUPE + BULLET SAFE)
# - Finds matching tables even when nested inside layout tables
# - Dedupes by fingerprint so nested traversal doesn't multiply tables
# - Keeps multi-line / bullet list cells (joins with "; ")
# - Skips instruction/rubric tables more safely
# - If stem not found, creates a sensible fallback from table headers
# ===================================================

MATCHING_STEM_RE = re.compile(
    r"\b("
    r"complete\s+the\s+table|"
    r"drag(?:ging)?\s+and\s+drop(?:ping)?|"
    r"drag\s+and\s+drop|"
    r"match\s+each|"
    r"match\s+the|"
    r"match\s+.*\s+to\s+the\s+correct|"
    r"select\s+one.*for\s+each"
    r")\b",
    re.IGNORECASE
)

INSTRUCTION_TABLE_NOISE_RE = re.compile(
    r"\b("
    r"range\s+and\s+conditions?|decision-?making\s+rules?|"
    r"rubric|pre-?approved\s+reasonable\s+adjustments?|"
    r"for\s+learners?|for\s+assessors?|instructions?|"
    r"evidence|required|criteria|competent|nyc|submission|marking"
    r")\b",
    re.IGNORECASE
)

def looks_like_matching_stem(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors")):
        return False
    if COMMAND_QUESTION_RE.match(t2):  # avoid essay stems stealing matching
        return False
    if "which of the following" in low:  # avoid mcq stems
        return False
    return bool(MATCHING_STEM_RE.search(t2))

def join_lines(lines: list[str]) -> str:
    parts = [clean_text(x) for x in (lines or []) if clean_text(x)]
    return "; ".join(parts).strip()

def cell_lines(cell) -> list[str]:
    """
    Bullet-safe cell extraction:
    - Each paragraph becomes a line (works even when Word bullets aren't '•')
    - If a paragraph contains '•' inline, we split it too.
    """
    lines = []
    for p in cell.paragraphs:
        t, _ = paragraph_text_and_is_red(p)
        t = clean_text(t)
        if not t:
            continue

        # If Word exported bullet symbol in same paragraph
        if "•" in t:
            parts = [x.strip() for x in t.split("•") if x.strip()]
            lines.extend(parts)
        else:
            lines.append(t)

    # de-dupe keep order
    out, seen = [], set()
    for x in lines:
        x = clean_text(x)
        if not x:
            continue
        k = normalize_key(x)
        if k in seen:
            continue
        seen.add(k)
        out.append(x)
    return out

def table_to_grid(table: Table) -> list[list[list[str]]]:
    return [[cell_lines(c) for c in row.cells] for row in table.rows]

def table_fingerprint(grid) -> str:
    flat = []
    for row in grid:
        for cell in row:
            flat.append("|".join(cell))
    blob = "||".join([normalize_key(x) for x in flat if x])
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()

def is_instruction_table(grid) -> bool:
    """
    Skip obvious policy/rubric tables.
    BUT don't be too aggressive (so we don't kill real matching).
    """
    texts = []
    for row in grid:
        for cell in row:
            if cell:
                texts.append(" ".join(cell))
    blob = " ".join(texts).strip()
    if not blob:
        return True

    # If the first row looks like a policy header table
    first_row = " ".join([join_lines(c) for c in (grid[0] if grid else [])]).lower()
    if "range and conditions" in first_row or "decision-making rules" in first_row or "pre-approved" in first_row:
        return True

    hits = sum(1 for t in texts if INSTRUCTION_TABLE_NOISE_RE.search(t))
    ratio = hits / max(1, len(texts))

    # only skip if *many* cells look like policy text
    return ratio >= 0.40

def guess_header_skip(grid) -> int:
    if not grid or not grid[0]:
        return 0

    # Join first row text
    row0 = " ".join([join_lines(c) for c in grid[0] if c]).strip().lower()

    # Common header words (add anything you see in your files)
    header_words = [
        "definition", "term", "meaning", "word", "concept",
        "numbers", "number", "example", "type", "classification",
        "left", "right"
    ]

    # If first row looks like headers, skip it
    if any(w in row0 for w in header_words):
        return 1

    # Also skip if it's mostly very short labels (typical header row)
    nonempty = [join_lines(c) for c in grid[0] if join_lines(c)]
    if nonempty and sum(len(x) <= 20 for x in nonempty) / len(nonempty) >= 0.8:
        return 1

    return 0


def pair_is_valid(left: str, right: str) -> bool:
    if not left or not right:
        return False
    if normalize_key(left) == normalize_key(right):
        return False
    # reject huge paragraphs (usually rubric/policy)
    if len(left) > 180 or len(right) > 350:
        return False
    return True

def score_columns(grid, a: int, b: int) -> int:
    start = guess_header_skip(grid)
    score = 0
    for r in range(start, len(grid)):
        if a >= len(grid[r]) or b >= len(grid[r]):
            continue
        left = join_lines(grid[r][a])
        right = join_lines(grid[r][b])
        if pair_is_valid(left, right):
            score += 1
    return score

def pick_best_columns(grid) -> tuple[int, int] | None:
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    if max_cols < 2:
        return None

    best, best_score = None, 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = score_columns(grid, a, b)
            if sc > best_score:
                best_score = sc
                best = (a, b)

    # require at least 2 usable pairs
        if best is None or best_score < 2:
            return None
        return best


def extract_pairs(grid, left_col: int, right_col: int, start_row: int = 0) -> list[dict]:
    pairs = []

    for r in range(start_row, len(grid)):
        if left_col >= len(grid[r]) or right_col >= len(grid[r]):
            continue

        left = join_lines(grid[r][left_col])
        right = join_lines(grid[r][right_col])

        left = re.sub(r"^\(?[a-z]\)\s*", "", left, flags=re.IGNORECASE).strip()
        left = re.sub(r"^[a-z]\.\s*", "", left, flags=re.IGNORECASE).strip()

        if not pair_is_valid(left, right):
            continue

        pairs.append({"left": left, "right": right})

    seen, out = set(), []
    for p in pairs:
        k = normalize_key(p["left"]) + "=>" + normalize_key(p["right"])
        if k in seen:
            continue
        seen.add(k)
        out.append(p)
    return out


def iter_elements_recursive(container):
    """
    Yield paragraphs + tables in document order, recursively inside tables.
    SAFE: no Document/DocxDocument isinstance checks.
    """
    # If it has `.element`, it's a Document-like object
    if hasattr(container, "element"):
        parent_elm = container.element.body
        parent_obj = container
    else:
        # table cell
        parent_elm = container._tc
        parent_obj = container

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", Paragraph(child, parent_obj))
        elif child.tag.endswith("}tbl"):
            tbl = Table(child, parent_obj)
            yield ("tbl", tbl)
            for row in tbl.rows:
                for cell in row.cells:
                    yield from iter_elements_recursive(cell)


def parse_matching_questions_doc_order(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    out = []
    recent_paras: list[str] = []
    seen_tables = set()
    MAX_LOOKBACK = 50

    def choose_stem() -> str | None:
        for t in reversed(recent_paras[-MAX_LOOKBACK:]):
            if looks_like_matching_stem(t):
                return strip_q_prefix(clean_text(t))
        return None

    for kind, el in iter_elements_recursive(doc):
        if kind == "p":
            t, _ = paragraph_text_and_is_red(el)
            t = clean_text(t)
            if t:
                recent_paras.append(t)
                if len(recent_paras) > 400:
                    recent_paras = recent_paras[-400:]
            continue

        # table
        grid = table_to_grid(el)
        header_skip = guess_header_skip_by_row_color(el)  # ✅ uses row shading

        tfp = table_fingerprint(grid)
        if tfp in seen_tables:
            continue
        seen_tables.add(tfp)

        if is_instruction_table(grid):
            continue

        # ✅ If one column is shaded (Term column), force it to be LEFT
        term_col = pick_term_column_by_fill(el)

        if term_col is not None:
    # choose RIGHT column by best score against the forced LEFT
            max_cols = max(len(r) for r in grid)
            best_right = None
            best_score = 0
            for b in range(max_cols):
                if b == term_col:
                    continue
                sc = score_columns(grid, term_col, b)
                if sc > best_score:
                    best_score = sc
                    best_right = b

            if best_right is None or best_score < 2:
                continue

            left_col, right_col = term_col, best_right

        else:
            cols = pick_best_columns(grid)
            if not cols:
                continue
            left_col, right_col = cols

        pairs = extract_pairs(grid, left_col, right_col, start_row=header_skip)

        if len(pairs) < 2:
            continue

        # prefer nearby stem, but DO NOT require it
        stem = choose_stem()
        if not stem:
            # fallback stem from header row if present
            header = grid[0] if grid else []
            hL = (header[left_col][0] if header and left_col < len(header) and header[left_col] else "Left")
            hR = (header[right_col][0] if header and right_col < len(header) and header[right_col] else "Right")
            stem = f"Match each '{hL}' to the correct '{hR}'."

        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False
        })

        # reset so the next table doesn't steal the same stem
        recent_paras = []

    return out


def cell_fill_hex(cell) -> str | None:
    """
    Returns background fill like 'D9D9D9' or None if no shading.
    """
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill:
        return None
    fill = fill.strip().upper()
    if fill in ("AUTO", "FFFFFF"):  # treat white as "no shading"
        return None
    return fill

def row_fill_signature(row_cells) -> tuple:
    # tuple of fills per cell so we can compare rows
    return tuple(cell_fill_hex(c) for c in row_cells)

def guess_header_skip_by_row_color(table: Table) -> int:
    """
    If row 0 is shaded (grey) and most body rows are not, skip it.
    """
    if not table.rows:
        return 0

    row0 = table.rows[0]
    sig0 = row_fill_signature(row0.cells)

    # If header row has ANY fill, and the next few rows mostly don't → header
    if any(sig0):
        sample = table.rows[1: min(6, len(table.rows))]
        non_header_like = 0
        for r in sample:
            sig = row_fill_signature(r.cells)
            if not any(sig):  # no fill
                non_header_like += 1
        if non_header_like >= max(1, len(sample) - 1):
            return 1

    return 0

def column_fill_stats(table: Table):
    """
    For each column index, count how many cells have non-empty fill.
    Returns: dict[col_index] = count
    """
    stats = {}
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            f = cell_fill_hex(cell)   # ✅ correct function name
            if f:
                stats[ci] = stats.get(ci, 0) + 1
    return stats



def pick_term_column_by_fill(table: Table) -> int | None:
    """
    If one column is clearly shaded more than others, treat it as the Term column.
    """
    stats = column_fill_stats(table)
    if not stats:
        return None

    # pick the column with most shaded cells
    best_col = max(stats, key=lambda c: stats[c])
    best = stats[best_col]

    # require it to be "clearly" more shaded than next best
    sorted_counts = sorted(stats.values(), reverse=True)
    second = sorted_counts[1] if len(sorted_counts) > 1 else 0

    # tune thresholds if needed
    if best >= 2 and best >= second + 2:
        return best_col
    return None

# ===================================================
# CANVAS API HELPERS
# ===================================================
def canvas_headers(canvas_token: str):
    return {"Authorization": f"Bearer {canvas_token}"}

def canvas_whoami(canvas_base_url: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/users/self"
    r = requests.get(url, headers=canvas_headers(canvas_token), timeout=30)
    if r.status_code == 401:
        return None
    r.raise_for_status()
    return r.json()

def list_courses(canvas_base_url: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/courses"
    out = []
    page = 1
    while True:
        r = requests.get(
            url,
            headers=canvas_headers(canvas_token),
            params={"per_page": 100, "page": page},
            timeout=60
        )
        r.raise_for_status()
        batch = r.json()
        if not batch:
            break
        out.extend(batch)
        page += 1
    return out

def get_existing_quiz_titles(canvas_base_url: str, course_id: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes"
    titles = set()
    page = 1
    while True:
        r = requests.get(
            url,
            headers=canvas_headers(canvas_token),
            params={"page": page, "per_page": 100},
            timeout=60
        )
        r.raise_for_status()
        data = r.json()
        if not data:
            break
        for q in data:
            titles.add((q.get("title") or "").strip())
        page += 1
    return titles

def generate_unique_title(base_title, existing_titles):
    if base_title not in existing_titles:
        return base_title
    i = 1
    while True:
        candidate = f"{base_title} ({i})"
        if candidate not in existing_titles:
            return candidate
        i += 1

def create_canvas_quiz(canvas_base_url: str, course_id: str, canvas_token: str,
                       title: str, description_html: str = "", settings: dict | None = None) -> int:
    settings = settings or {}

    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes"
    quiz_obj = {
        "title": title,
        "description": description_html,
        "published": False,
        "quiz_type": "assignment",
        "shuffle_answers": bool(settings.get("shuffle_answers", False)),
        "one_question_at_a_time": bool(settings.get("one_question_at_a_time", False)),
        "show_correct_answers": bool(settings.get("show_correct_answers", False)),
        "scoring_policy": settings.get("scoring_policy", "keep_highest"),
    }

    tl = int(settings.get("time_limit", 0) or 0)
    if tl > 0:
        quiz_obj["time_limit"] = tl

    allow_multi = bool(settings.get("allow_multiple_attempts", False))
    if allow_multi:
        aa = int(settings.get("allowed_attempts", 2) or 2)
        quiz_obj["allowed_attempts"] = max(2, aa)

    if bool(settings.get("access_code_enabled", False)) and (settings.get("access_code") or "").strip():
        quiz_obj["access_code"] = settings["access_code"].strip()

    for k in ["due_at", "unlock_at", "lock_at"]:
        v = (settings.get(k) or "").strip()
        if v:
            quiz_obj[k] = v

    payload = {"quiz": quiz_obj}

    r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    if r.status_code == 401:
        raise RuntimeError("401 Unauthorized — token invalid/expired.")
    if r.status_code == 403:
        raise RuntimeError("403 Forbidden — missing permission in this course.")
    r.raise_for_status()
    return r.json()["id"]

def publish_quiz(canvas_base_url: str, course_id: str, canvas_token: str, quiz_id: int):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes/{quiz_id}"
    payload = {"quiz": {"published": True}}
    r = requests.put(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    r.raise_for_status()

def add_question_to_quiz(canvas_base_url: str, course_id: str, canvas_token: str, quiz_id: int, q: dict):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes/{quiz_id}/questions"
    qtext = strip_q_prefix((q.get("question") or "").strip())
    kind = (q.get("kind") or "").lower()

    # MATCHING
    if kind == "matching":
        pairs = q.get("pairs") or []
        answers = []
        for p in pairs:
            left = (p.get("left") or "").strip()
            right = (p.get("right") or "").strip()
            if left and right:
                answers.append({
                    "answer_match_left": left,
                    "answer_match_right": right,
                    "answer_weight": 100
                })

        payload = {
            "question": {
                "question_name": qtext[:100] if qtext else "Matching",
                "question_text": qtext,
                "question_type": "matching_question",
                "points_possible": 1,
                "answers": answers
            }
        }
        r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
        if r.status_code >= 400:
            raise RuntimeError(f"Canvas error {r.status_code}: {r.text[:600]}")
        r.raise_for_status()
        return

    # ESSAY
    opts = [o.strip() for o in (q.get("options") or []) if o and o.strip()]
    correct = q.get("correct", []) or []
    if kind == "essay" or len(opts) < 2:
        payload = {
            "question": {
                "question_name": qtext[:100] if qtext else "Question",
                "question_text": qtext or " ",
                "question_type": "essay_question",
                "points_possible": 1
            }
        }
        r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
        r.raise_for_status()
        return

    # MCQ
    qlower = (qtext or "").lower()
    multi = bool(q.get("multi")) or (len(correct) > 1) or bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower)) or ("apply" in qlower)
    qtype = "multiple_answers_question" if multi else "multiple_choice_question"

    answers = [{"answer_text": opt, "answer_weight": 100 if idx in correct else 0} for idx, opt in enumerate(opts)]

    payload = {
        "question": {
            "question_name": (qtext[:100] if qtext else "Question"),
            "question_text": qtext,
            "question_type": qtype,
            "points_possible": 1,
            "answers": answers
        }
    }

    r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    if r.status_code >= 400:
        raise RuntimeError(f"Canvas error {r.status_code}: {r.text[:600]}")
    r.raise_for_status()

# ===================================================
# SIDEBAR: LOGIN + COURSE PICKER
# ===================================================
with st.sidebar:
    st.header("🔐 Canvas Login (Token)")

    st.session_state.canvas_base_url = st.text_input(
        "Canvas Base URL",
        value=st.session_state.canvas_base_url,
        help="Example: https://learningvault.test.instructure.com/api/v1"
    ).strip()

    st.session_state.canvas_token = st.text_input(
        "Canvas Access Token",
        value=st.session_state.canvas_token,
        type="password",
        help="Canvas API uses access tokens."
    )

    c_login, c_logout = st.columns(2)

    if c_login.button("Login"):
        try:
            me = canvas_whoami(st.session_state.canvas_base_url, st.session_state.canvas_token)
            if me:
                st.session_state.logged_in = True
                st.session_state.me = me
                st.session_state.courses_cache = None
                st.success(f"Logged in as: {me.get('name', 'Unknown')}")
            else:
                st.session_state.logged_in = False
                st.session_state.me = None
                st.error("Login failed: token invalid/expired.")
        except Exception as e:
            st.session_state.logged_in = False
            st.session_state.me = None
            st.error(f"Login failed: {e}")

    if c_logout.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.me = None
        st.session_state.selected_course_id = None
        st.session_state.courses_cache = None
        st.session_state.questions = []
        st.session_state.parsed_ok = False
        st.info("Logged out.")

    st.divider()

    if st.session_state.logged_in:
        st.subheader("✅ Course")
        try:
            if st.session_state.courses_cache is None:
                st.session_state.courses_cache = list_courses(
                    st.session_state.canvas_base_url,
                    st.session_state.canvas_token
                )

            courses = st.session_state.courses_cache or []
            if not courses:
                st.warning("No courses visible to this token.")
            else:
                label_to_id = {}
                labels = []
                for c in courses:
                    cid = c.get("id")
                    name = (c.get("name") or c.get("course_code") or f"Course {cid}").strip()
                    label = f"{name} (ID: {cid})"
                    labels.append(label)
                    label_to_id[label] = str(cid)

                default_index = 0
                if st.session_state.selected_course_id:
                    for i, lb in enumerate(labels):
                        if label_to_id[lb] == st.session_state.selected_course_id:
                            default_index = i
                            break

                chosen = st.selectbox("Select course", labels, index=default_index)
                st.session_state.selected_course_id = label_to_id[chosen]

        except Exception as e:
            st.error(f"Failed to load courses: {e}")
    else:
        st.info("Login first to select course and upload.")

# Block main area until logged in + course chosen
if not st.session_state.logged_in:
    st.warning("Please login in the sidebar first.")
    st.stop()

if not st.session_state.selected_course_id:
    st.warning("Please select a course in the sidebar.")
    st.stop()

course_id = st.session_state.selected_course_id
canvas_base_url = st.session_state.canvas_base_url
canvas_token = st.session_state.canvas_token

# ===================================================
# MAIN UI: UPLOAD + PARSE
# ===================================================
st.subheader("1) Upload DOCX and Parse")

uploaded = st.file_uploader("📤 Upload DOCX assessment file", type=["docx"])
colA, colB = st.columns([1, 1])
parse_btn = colA.button("🧠 Parse DOCX", type="primary", disabled=(uploaded is None))
reset_btn = colB.button("♻️ Reset parsed data")

log_box = st.empty()

if reset_btn:
    st.session_state.docx_filename = None
    st.session_state.description_html = ""
    st.session_state.questions = []
    st.session_state.parsed_ok = False
    st.success("Reset complete.")

if parse_btn:
    if not uploaded:
        st.warning("Upload a DOCX first.")
        st.stop()

    # ✅ clear previous results FIRST (prevents “previous answers carry over”)
    st.session_state.questions = []
    st.session_state.parsed_ok = False
    st.session_state.description_html = ""
    st.session_state.docx_filename = None
        # ✅ bump parse run id so Streamlit forgets old widget values
    st.session_state.parse_run_id += 1


    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(uploaded.read())
        docx_path = tmp.name

    buf = io.StringIO()
    try:
        with contextlib.redirect_stdout(buf):
            items = extract_items_with_red(docx_path)
            description_html = build_description(items)

            matching = parse_matching_questions_doc_order(docx_path)
            mcq = parse_mcq_questions(items)
            essay = parse_essay_questions(items)

            questions = matching + mcq + essay
            questions = dedupe_questions(matching+mcq+essay)

            print("DEBUG: items extracted:", len(items))
            print("DEBUG: matching:", len(matching), "mcq:", len(mcq), "essay:", len(essay))
            print("Parsed questions:", len(questions))

            if not questions:
                raise RuntimeError("No questions detected. Check: headings, question wording, or doc format.")

        st.session_state.docx_filename = uploaded.name
        st.session_state.description_html = description_html
        st.session_state.questions = questions
        st.session_state.parsed_ok = True

        log_box.code(buf.getvalue())
        st.success(f"✅ Parsed {len(questions)} questions.")

    except Exception as e:
        log_box.code(buf.getvalue())
        st.error(f"❌ Parse error: {e}")

    finally:
        try:
            os.remove(docx_path)
        except Exception:
            pass

# ===================================================
# STOP if nothing parsed yet
# ===================================================
if not st.session_state.parsed_ok:
    st.info("Parse a DOCX to continue (you can still select course and login first).")
    st.stop()

questions = st.session_state.questions

st.divider()

# ===================================================
# DETAILS
# ===================================================
st.subheader("2) Details (Canvas Quiz Settings)")

default_title = os.path.splitext(st.session_state.docx_filename or "Quiz")[0]
quiz_title = st.text_input("Quiz Title *", value=default_title)

quiz_instructions = st.text_area(
    "Quiz Instructions (HTML allowed)",
    value=st.session_state.description_html or "",
    height=160
)

d = st.session_state.details

c1, c2, c3 = st.columns(3)
d["shuffle_answers"] = c1.checkbox("Shuffle Answers", value=bool(d.get("shuffle_answers", True)))
d["one_question_at_a_time"] = c2.checkbox("Show one question at a time", value=bool(d.get("one_question_at_a_time", False)))
d["show_correct_answers"] = c3.checkbox("Let Students See The Correct Answers", value=bool(d.get("show_correct_answers", False)))

c4, c5, c6 = st.columns(3)
d["time_limit"] = c4.number_input("Time Limit (minutes, 0 = none)", min_value=0, value=int(d.get("time_limit", 0)), step=5)
d["allow_multiple_attempts"] = c5.checkbox("Allow Multiple Attempts", value=bool(d.get("allow_multiple_attempts", False)))

if d["allow_multiple_attempts"]:
    current_attempts = int(d.get("allowed_attempts", 2) or 2)
    if current_attempts < 2:
        current_attempts = 2
        d["allowed_attempts"] = 2
    d["allowed_attempts"] = c5.number_input("Allowed Attempts", min_value=2, value=current_attempts, step=1)
else:
    d["allowed_attempts"] = 1

d["scoring_policy"] = c6.selectbox(
    "Quiz Score to Keep",
    ["keep_highest", "keep_latest"],
    index=0 if d.get("scoring_policy", "keep_highest") == "keep_highest" else 1
)

st.markdown("**Quiz Restrictions**")
d["access_code_enabled"] = st.checkbox("Require an access code", value=bool(d.get("access_code_enabled", False)))
if d["access_code_enabled"]:
    d["access_code"] = st.text_input("Access code", value=(d.get("access_code") or ""))
else:
    d["access_code"] = ""

st.markdown("**Assign / Availability (optional, ISO datetime)**")
st.caption("Example: 2026-01-20T23:59:00Z (If you don’t know, leave blank for now.)")

cc1, cc2, cc3 = st.columns(3)
d["due_at"] = cc1.text_input("Due Date (due_at)", value=(d.get("due_at") or ""))
d["unlock_at"] = cc2.text_input("Available from (unlock_at)", value=(d.get("unlock_at") or ""))
d["lock_at"] = cc3.text_input("Until (lock_at)", value=(d.get("lock_at") or ""))

st.session_state.details = d

st.divider()

# ===================================================
# QUESTIONS EDITOR + PAGING
# ===================================================
st.subheader("3) Questions")

colp1, colp2 = st.columns([1, 1])
page_size = colp1.selectbox("Questions per page", [5, 10, 15, 20, 30], index=1)
total = len(questions)
total_pages = max(1, math.ceil(total / page_size))
page = colp2.number_input("Page", min_value=1, max_value=total_pages, value=1, step=1)

start = (page - 1) * page_size
end = min(start + page_size, total)
st.caption(f"Showing questions {start+1}–{end} of {total}")

edited = [q.copy() for q in questions]

for i in range(start, end):
    q = edited[i]
    run = st.session_state.parse_run_id

    kind = (q.get("kind") or "").lower()
    preview = strip_q_prefix(q.get("question", ""))[:90]
    label_kind = "Matching" if kind == "matching" else ("Essay/Short Answer" if kind == "essay" else "MCQ")

    with st.expander(f"Q{i+1} ({label_kind}): {preview}"):
        q_text = st.text_area(
            "Question text",
            value=q.get("question", ""),
            key=f"{run}_qtext_{i}",
            height=90
        )

        q["question"] = strip_q_prefix(q_text.strip())

        if kind == "essay":
            st.info("This question will be uploaded as an ESSAY (student types the answer).")
            q["options"] = []
            q["correct"] = []
            q["multi"] = False

        elif kind == "matching":
            st.info("This question will be uploaded as MATCHING (left item → dropdown right item).")
            st.caption("Tip: if right side was a bullet list in Word, it appears joined with '; ' — that is correct.")
            pairs = q.get("pairs") or []
            new_pairs = []
            for j, p in enumerate(pairs):
                lc1, lc2 = st.columns([0.6, 0.4])
                left = lc1.text_input(
                    f"Left (row {j+1})",
                    value=p.get("left",""),
                    key=f"{run}_match_{i}_l_{j}"
                )
                right = lc2.text_input(
                    f"Right (row {j+1})",
                    value=p.get("right",""),
                    key=f"{run}_match_{i}_r_{j}"
                )

                if left.strip() and right.strip():
                    new_pairs.append({"left": left.strip(), "right": right.strip()})
            q["pairs"] = new_pairs

        else:
            opts = q.get("options", []) or []
            correct_set = set(q.get("correct", []) or [])

            st.write("**Options** (tick ✅ for correct answer)")
            new_opts = []
            new_correct = []

            for j, opt in enumerate(opts):
                oc1, oc2 = st.columns([0.12, 0.88])
                is_corr = oc1.checkbox(
                    "",
                    value=(j in correct_set),
                    key=f"{run}_q{i}_corr_{j}"
                )
                opt_text = oc2.text_input(
                    f"Option {j+1}",
                    value=opt,
                    key=f"{run}_q{i}_opt_{j}"
                )

                new_opts.append(opt_text.strip())
                if is_corr:
                    new_correct.append(j)

            add_opt = st.text_input(
                "New option text (optional)",
                value="",
                key=f"{run}_q{i}_newopt"
            )

            if add_opt.strip():
                new_opts.append(add_opt.strip())

            cleaned_opts = []
            idx_map = {}
            for old_index, txt in enumerate(new_opts):
                if txt.strip():
                    idx_map[old_index] = len(cleaned_opts)
                    cleaned_opts.append(txt.strip())

            remapped_correct = []
            for old_i in new_correct:
                if old_i in idx_map:
                    remapped_correct.append(idx_map[old_i])

            q["options"] = cleaned_opts
            q["correct"] = sorted(set(remapped_correct))

            qlower = (q.get("question") or "").lower()
            q["multi"] = (
                ("select" in qlower and any(w in qlower for w in ["two", "three", "four", "five"]))
                or ("apply" in qlower)
                or (len(q["correct"]) > 1)
            )

st.session_state.questions = edited

st.divider()

# ===================================================
# SAVE ACTIONS
# ===================================================
st.subheader("4) Save to Canvas")

colS1, colS2 = st.columns([1, 1])
save_draft = colS1.button("💾 Save to Canvas (Draft)")
save_publish = colS2.button("🚀 Save & Publish")

def validate_before_upload(qs: list[dict]) -> list[str]:
    problems = []
    for i, q in enumerate(qs, start=1):
        kind = (q.get("kind") or "").lower()
        qt = (q.get("question") or "").strip()

        if len(qt) < 10:
            problems.append(f"Q{i}: question text too short.")

        if kind == "matching":
            pairs = q.get("pairs") or []
            if len(pairs) < 2:
                problems.append(f"Q{i}: matching needs at least 2 pairs.")
            continue

        if kind == "essay":
            continue

        opts = q.get("options") or []
        corr = q.get("correct") or []
        if len(opts) >= 2 and len(corr) == 0:
            problems.append(f"Q{i}: no correct answer selected (red not detected or tick ✅).")
    return problems

if save_draft or save_publish:
    qs = st.session_state.questions
    probs = validate_before_upload(qs)
    if probs:
        st.error("Please fix these issues before uploading:")
        for p in probs[:15]:
            st.write(f"- {p}")
        st.stop()

    base_title = (quiz_title or "").strip() or default_title
    try:
        existing_titles = get_existing_quiz_titles(canvas_base_url, course_id, canvas_token)
        final_title = generate_unique_title(base_title, existing_titles)

        with st.spinner("Creating quiz in Canvas..."):
            quiz_id = create_canvas_quiz(
                canvas_base_url=canvas_base_url,
                course_id=course_id,
                canvas_token=canvas_token,
                title=final_title,
                description_html=quiz_instructions,
                settings=st.session_state.details
            )

        with st.spinner("Uploading questions..."):
            for q in qs:
                add_question_to_quiz(canvas_base_url, course_id, canvas_token, quiz_id, q)

        if save_publish:
            with st.spinner("Publishing quiz..."):
                publish_quiz(canvas_base_url, course_id, canvas_token, quiz_id)

        st.success("✅ Done!")
        st.write(f"**Quiz title:** {final_title}")
        st.write(f"**Quiz ID:** {quiz_id}")
        st.write(f"**Course ID:** {course_id}")
        st.info("Quiz published ✅" if save_publish else "Quiz saved as draft (unpublished).")

    except Exception as e:
        st.error(f"❌ Upload failed: {e}")

st.caption("Token login only (Canvas API does not support username/password).")
