# ==========================================================
# Learning Vault – Canvas Quiz Uploader (FULL STABLE VERSION)
# ✅ Matching + MCQ + Essay
# ✅ Matching fixed: no duplicates, correct stems, list-cells, no “wrong table”
# ✅ No undefined functions / no duplicate parsers
# ✅ Clears previous file results properly (no “previous answers” carry over)
# ==========================================================

import os
import re
import io
import math
import hashlib
import tempfile
import contextlib
import zipfile
import shutil
import traceback
import xml.etree.ElementTree as ET
from datetime import datetime, date, time

import pytz
import requests
import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P

# ===================================================
# TIMEZONE
# ===================================================
TZ_NAME = "Australia/Sydney"
tz = pytz.timezone(TZ_NAME)

def combine_date_time(d: date | None, t: time | None):
    if not d or not t:
        return None
    dt = tz.localize(datetime.combine(d, t))
    return dt.isoformat()

# ===================================================
# UI
# ===================================================
st.set_page_config(page_title="Learning Vault – Canvas Quiz Uploader", layout="wide")
st.title("Learning Vault – Canvas Quiz Uploader (Web Interface)")

# ===================================================
# SESSION STATE DEFAULTS
# ===================================================
def ss_init(key, value):
    if key not in st.session_state:
        st.session_state[key] = value

ss_init("logged_in", False)
ss_init("me", None)
ss_init("canvas_token", "")
ss_init("canvas_base_url", "https://learningvault.test.instructure.com/api/v1")
ss_init("courses_cache", None)
ss_init("selected_course_id", None)

ss_init("docx_filename", None)
ss_init("description_html", "")
ss_init("questions", [])
ss_init("parsed_ok", False)
ss_init("parse_run_id", 0)
ss_init("missing_phrase_debug", [])


ss_init("details", {
    "shuffle_answers": True,
    "time_limit": 0,
    "allow_multiple_attempts": False,
    "allowed_attempts": 2,
    "scoring_policy": "keep_highest",
    "one_question_at_a_time": False,
    "show_correct_answers": False,
    "access_code_enabled": False,
    "access_code": "",
    "due_at": "",
    "unlock_at": "",
    "lock_at": "",
})

# ===================================================
# HELPERS
# ===================================================
def clean_text(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "")).strip()

def normalize_key(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s

Q_PREFIX_RE = re.compile(
    r"^[^A-Za-z0-9]*(?:lo\s*)?(?:question|q)\s*\d+\s*(?:[\.\)]\s*|[:\-–—]\s+|\s+)",
    re.IGNORECASE
)
NUM_PREFIX_RE = re.compile(
    r"^[^A-Za-z0-9]*\(?\d+\)?\s*(?:[\.\)]\s*|[:\-–—]\s+)",
    re.IGNORECASE
)
def strip_q_prefix(line: str) -> str:
    s = (line or "").strip()
    s = Q_PREFIX_RE.sub("", s)
    s = NUM_PREFIX_RE.sub("", s)
    return s.strip()

LETTERED_OPT_PREFIX_RE = re.compile(r"^\s*(?:[\(\[]?[a-hA-H][\)\].:-])\s+")
def strip_lettered_prefix(t: str) -> str:
    return LETTERED_OPT_PREFIX_RE.sub("", (t or "").strip()).strip()

def question_fingerprint(q: dict) -> str:
    qt = normalize_key(q.get("question", ""))
    kind = normalize_key(q.get("kind", ""))
    opts = [normalize_key(x) for x in (q.get("options") or [])]
    pairs = q.get("pairs") or []
    pairs_blob = "||".join([normalize_key(p.get("left","")) + "=>" + normalize_key(p.get("right","")) for p in pairs])
    blob = kind + "||" + qt + "||" + "||".join(opts) + "||" + pairs_blob
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()

def dedupe_questions(questions: list) -> list:
    seen = set()
    out = []
    for q in questions:
        fp = question_fingerprint(q)
        if fp in seen:
            continue
        seen.add(fp)
        out.append(q)
    return out

def collapse_duplicate_mcq(questions: list[dict]) -> list[dict]:
    """
    Second-pass de-dupe for MCQ where parsing glitches can create near-duplicates:
    - same question text
    - different (often bloated) option lists due to stem-detection misses
    Keep the "best" candidate and preserve the earliest document order.
    """
    groups: dict[str, list[dict]] = {}
    non_mcq: list[dict] = []

    for q in questions:
        if (q.get("kind") or "").lower() == "mcq":
            key = normalize_key(q.get("question", ""))
            groups.setdefault(key, []).append(q)
        else:
            non_mcq.append(q)

    def score_mcq(q: dict) -> tuple[int, int]:
        opts = q.get("options") or []
        n = len(opts)
        score = 0

        if 2 <= n <= 6:
            score += 6
        elif 2 <= n <= 10:
            score += 3
        else:
            score -= 3

        if n > 12:
            score -= 8

        bad = 0
        for o in opts:
            low = normalize_key(o)
            if "?" in (o or ""):
                bad += 2
            if re.match(r"^(?:in\\s+regard\\b|(?:in\\s+)?which\\b|what\\b|why\\b|how\\b|select\\b|choose\\b|pick\\b)", low):
                bad += 3
        score -= bad

        return (score, -n)

    out = list(non_mcq)
    for _, lst in groups.items():
        if len(lst) == 1:
            out.append(lst[0])
            continue

        best = max(lst, key=score_mcq)
        best_order = min(int(x.get("_order", 10**9)) for x in lst)
        best["_order"] = best_order
        out.append(best)

    out.sort(key=lambda q: int(q.get("_order", 10**9)))
    return out

def docx_xml_contains(docx_path: str, phrase: str) -> bool:
    """
    Heuristic check: if a string isn't present in any DOCX text nodes, it's likely embedded as an image.
    We search actual XML text nodes (e.g. w:t, a:t) to avoid false negatives from tag-splitting.
    """
    if not phrase:
        return False
    needle = re.sub(r"\s+", " ", phrase).strip().lower()
    try:
        with zipfile.ZipFile(docx_path) as z:
            for name in z.namelist():
                if not name.lower().endswith(".xml"):
                    continue
                try:
                    raw = z.read(name)
                except Exception:
                    continue
                try:
                    root = ET.fromstring(raw)
                except Exception:
                    # Fallback: raw substring scan (less reliable, but better than nothing)
                    try:
                        txt = raw.decode("utf-8", errors="ignore")
                    except Exception:
                        continue
                    if needle in re.sub(r"\s+", " ", txt).lower():
                        return True
                    continue

                parts: list[str] = []
                for el in root.iter():
                    tag = str(getattr(el, "tag", ""))
                    local = tag.rsplit("}", 1)[-1]
                    if local != "t":
                        continue
                    txt = getattr(el, "text", None)
                    if txt:
                        parts.append(txt)

                blob = re.sub(r"\s+", " ", "".join(parts)).strip().lower()
                if needle and needle in blob:
                    return True
    except Exception:
        return False
    return False

def save_docx_debug_copy(src_docx_path: str, original_name: str | None = None) -> str | None:
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        if os.path.basename(base_dir) == "backups":
            base_dir = os.path.dirname(base_dir)
        uploads_dir = os.path.join(base_dir, "uploads")
        os.makedirs(uploads_dir, exist_ok=True)
        base = os.path.basename(original_name or "debug.docx")
        base = re.sub(r"[^A-Za-z0-9._ -]+", "_", base).strip() or "debug.docx"
        out_path = os.path.join(uploads_dir, f"debug_{base}")
        shutil.copyfile(src_docx_path, out_path)
        return out_path
    except Exception:
        return None

def repair_placeholder_mcq_options(mcq: list[dict], items: list[dict], item_index: dict[str, int]) -> list[dict]:
    """
    If an MCQ has placeholder options, try to recover options by scanning forward in `items`.
    This fixes cases where the MCQ parser fails to carry over lettered '(a) ...' option lines.
    """
    if not mcq:
        return mcq

    question_start_re = re.compile(r"^(?:in\s+)?(which|what|why|how)\b", re.IGNORECASE)
    select_stem_re = re.compile(
        r"^(?:q\s*\d+\.?\s*)?(select|choose|pick)\s+the\s+(best|correct|most\s+appropriate)\b",
        re.IGNORECASE
    )
    contains_select_summary_re = re.compile(
        r"\b(select|choose|pick)\s+the\s+(best|correct|most\s+appropriate)\s+summary\b",
        re.IGNORECASE
    )
    best_match_re = re.compile(
        r"\b(best\s+match|does\s+the\s+following\s+description\s+best\s+match)\b",
        re.IGNORECASE
    )

    def is_placeholder_opt(opt: str) -> bool:
        return (opt or "").startswith("⚠ Option text not extracted")

    def looks_like_new_stem(line: str) -> bool:
        t = strip_q_prefix(clean_text(line))
        if not t:
            return False
        if looks_like_matching_stem(t):
            return True
        if COMMAND_QUESTION_RE.match(t):
            return True
        if QUESTION_CMD_INNER_RE.search(t):
            return True
        if select_stem_re.match(t):
            return True
        if contains_select_summary_re.search(t) or best_match_re.search(t):
            return True
        if question_start_re.match(t) and "?" in t:
            return True
        return False

    def looks_like_option_line(line: str, is_red: bool) -> bool:
        t = strip_q_prefix(clean_text(line))
        if not t:
            return False
        if STOP_OPTION_RE.match(t) or NOISE_RE.match(t) or OPTION_NOISE_RE.match(t):
            return False
        if looks_like_new_stem(t):
            return False
        if LETTERED_OPT_PREFIX_RE.match(t):
            return True
        if is_red and len(t) <= 320 and not t.endswith("?"):
            return True
        return False

    out = []
    for q in mcq:
        opts = q.get("options") or []
        if len(opts) >= 2 and all(is_placeholder_opt(o) for o in opts):
            key = normalize_key(q.get("question", ""))
            idx = item_index.get(key)
            if idx is None:
                base = re.sub(r"\s*\((?:select|choose).*\)\s*$", "", (q.get("question") or ""), flags=re.IGNORECASE).strip()
                idx = item_index.get(normalize_key(base))

            if idx is not None:
                recovered = []
                recovered_red = []
                for j in range(idx + 1, min(len(items), idx + 60)):
                    t = items[j].get("text", "")
                    is_red = bool(items[j].get("is_red", False))
                    if STOP_OPTION_RE.match(t):
                        break
                    if looks_like_new_stem(t):
                        break
                    if looks_like_option_line(t, is_red):
                        recovered.append(strip_lettered_prefix(t))
                        recovered_red.append(is_red)

                final_opts = []
                final_corr = []
                seen = set()
                for t, is_red in zip(recovered, recovered_red):
                    t = clean_text(t)
                    if not t:
                        continue
                    nk = normalize_key(t)
                    if nk in seen:
                        continue
                    seen.add(nk)
                    final_opts.append(t)
                    if is_red:
                        final_corr.append(len(final_opts) - 1)

                if len(final_opts) >= 2:
                    q = q.copy()
                    q["options"] = final_opts
                    q["correct"] = final_corr
                    q["multi"] = bool(q.get("multi")) or (len(final_corr) > 1) or ("apply" in (q.get("question") or "").lower())
        out.append(q)

    return out

# ===================================================
# RED DETECTION (correct answers in red)
# ===================================================
def is_red_run(run) -> bool:
    color = run.font.color
    if not color:
        return False
    rgb = color.rgb
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return (r >= 200 and g <= 80 and b <= 80)

def paragraph_text_and_is_red(paragraph):
    text = "".join(run.text for run in paragraph.runs).strip()
    any_red = any(is_red_run(run) and run.text.strip() for run in paragraph.runs)
    return text, any_red

def is_red_hex(val: str) -> bool:
    v = (val or "").strip().lstrip("#").upper()
    if not re.fullmatch(r"[0-9A-F]{6}", v):
        return False
    r, g, b = int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16)
    return (r >= 200 and g <= 80 and b <= 80)

def txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    """
    Extract text from a raw WordprocessingML paragraph (used for text boxes).
    """
    parts: list[str] = []
    any_red = False

    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue

        run_texts = []
        for t_node in r.iter():
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None):
                run_texts.append(t_node.text)
        t = "".join(run_texts)
        if t:
            parts.append(t)

        if any_red or not t.strip():
            continue

        rpr = None
        for child in list(r):
            if str(getattr(child, "tag", "")).endswith("}rPr"):
                rpr = child
                break

        if rpr is None:
            continue

        color = None
        for c in rpr.iter():
            if str(getattr(c, "tag", "")).endswith("}color"):
                color = c
                break

        if color is None:
            continue

        val = None
        for k, v in getattr(color, "attrib", {}).items():
            if str(k).endswith("}val") or str(k) == "val":
                val = v
                break
        if val and is_red_hex(val):
            any_red = True

    return "".join(parts).strip(), any_red

def textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    """
    python-docx doesn't surface text inside drawing text boxes (w:txbxContent).
    This extracts any such text anchored in the given paragraph.
    """
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    out: list[tuple[str, bool]] = []

    txbx_paras = []
    for el in p_elm.iter():
        if str(getattr(el, "tag", "")).endswith("}txbxContent"):
            for p2 in el.iter():
                if str(getattr(p2, "tag", "")).endswith("}p"):
                    txbx_paras.append(p2)

    for tx_p in txbx_paras:
        t, red = txbx_paragraph_text_and_is_red(tx_p)
        t = clean_text(t)
        if t:
            out.append((t, red))
    return out

def drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    """
    Extract text from DrawingML shapes (e.g. black boxes with red text) which often
    live under a:txBody/a:p/a:r/a:t rather than w:txbxContent.
    """
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    lines: list[tuple[str, bool]] = []
    seen = set()

    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue

        parts: list[str] = []
        any_red = False

        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = None
                for k, v in getattr(node, "attrib", {}).items():
                    if str(k).endswith("}val") or str(k) == "val":
                        val = v
                        break
                if val and is_red_hex(val):
                    any_red = True

        text = clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))

    return lines

# ===================================================
# DOCX ORDER ITERATOR (TOP LEVEL ONLY)
# Important: no nested recursion here → stops “2 matching became 6”
# ===================================================
def iter_block_items(doc):
    """
    Yield Paragraph and Table objects in true top-level document order.
    SAFE: does not rely on isinstance(Document).
    """
    body = doc.element.body

    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


# ===================================================
# EXTRACT ITEMS (paragraph + table cells in true top-level order)
# for MCQ/Essay parsing + description block
# ===================================================
def extract_items_with_red(docx_path):
    doc = Document(docx_path)
    items = []
    last_pushed: tuple[str, bool] | None = None

    def push_text(t: str, is_red: bool):
        nonlocal last_pushed
        t = clean_text(t)
        if t and (last_pushed != (t, is_red)):
            items.append({"text": t, "is_red": is_red})
            last_pushed = (t, is_red)

    def push_paragraph(p: Paragraph):
        t, red = paragraph_text_and_is_red(p)
        push_text(t, red)
        for t2, red2 in textbox_texts_in_paragraph(p):
            push_text(t2, red2)
        for t2, red2 in drawingml_texts_in_paragraph(p):
            push_text(t2, red2)

    def push_table(tbl: Table):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    push_paragraph(p)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            push_paragraph(block)
        else:
            push_table(block)

    # Some assessments put key text inside headers/footers (often as shapes/text boxes).
    # python-docx doesn't surface those shape texts, so we run our custom extractors there too.
    for sec in getattr(doc, "sections", []) or []:
        for hf in [
            getattr(sec, "header", None),
            getattr(sec, "footer", None),
            getattr(sec, "first_page_header", None),
            getattr(sec, "first_page_footer", None),
            getattr(sec, "even_page_header", None),
            getattr(sec, "even_page_footer", None),
        ]:
            if hf is None:
                continue
            for p in getattr(hf, "paragraphs", []) or []:
                push_paragraph(p)
            for tbl in getattr(hf, "tables", []) or []:
                push_table(tbl)

    return items

# ===================================================
# QUESTION / NOISE RULES
# ===================================================
NOISE_RE = re.compile(
    r"^(Instructions|For learners|For students|For assessors|Range and conditions|Decision-making rules|"
    r"Pre-approved reasonable adjustments|Rubric|Knowledge Test|"
    r"A rubric has been assigned\b|Answers will be assessed against\b|As a principle\b)\b",
    re.IGNORECASE
)

STOP_OPTION_RE = re.compile(
    r"^(Learner feedback|Assessment outcome|Assessor signature|Assessor name|Final comments)\b"
    r"|^(Competent|Not Yet Competent|NYC|C|Date)\s*[:\-]?$",
    re.IGNORECASE
)

OPTION_NOISE_RE = re.compile(
    r"^(Learning\s+Vault|\d{1,2}/\d{1,2}/\d{2,4}|SIT[A-Z0-9]{5,}\b)",
    re.IGNORECASE
)

QUESTION_CMD_INNER_RE = re.compile(
    r"\b(Which\s+of\s+the\s+following\b|"
    r"(Identify|Select|Choose|Pick)\s+(?:the\s+)?(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b)",
    re.IGNORECASE
)

COMMAND_QUESTION_RE = re.compile(
    r"^(Illustrate|Explain|Describe|Discuss|Outline|Compare|Summari[sz]e|"
    r"Critically\s+(?:assess|analyse|analyze|evaluate)|"
    r"Evaluate|Determine|Articulate|Prescribe|Analyse|Analyze|Review|Recommend|Provide)\b.+",
    re.IGNORECASE
)

RUBRIC_START_RE = re.compile(r"^Answer\s+needs\s+to\s+address\b", re.IGNORECASE)
ESSAY_GUIDE_RE = re.compile(r"^Answer\s+(may|must)\s+address", re.IGNORECASE)

# ===================================================
# ITEM REPAIR (merge split question stems)
# ===================================================
_DANGLING_Q_END_RE = re.compile(r"\b(of|for|to|with|and|or|in|on|at|from|by|as|about)\s*$", re.IGNORECASE)
def merge_dangling_question_lines(items: list[dict]) -> list[dict]:
    """
    Some DOCX exports split a question stem across two lines where the first line looks like:
      "Which ... requirements of"
    and the next line starts with a lowercase continuation, sometimes with a "Q15 ..." prefix.
    If we don't merge them, the first line can be mis-parsed as an option.
    """
    out: list[dict] = []
    i = 0
    n = len(items)

    while i < n:
        it = items[i]
        t = clean_text(it.get("text", ""))
        if not t:
            i += 1
            continue

        t_stem = strip_q_prefix(t)
        can_start_q = bool(re.match(r"^(?:in\s+)?(which|what|why|how)\b", t_stem, re.IGNORECASE))
        dangling = can_start_q and ("?" not in t_stem) and _DANGLING_Q_END_RE.search(t_stem)
        dangling = bool(dangling) and not looks_like_matching_stem(t_stem) and not NOISE_RE.match(t_stem) and not STOP_OPTION_RE.match(t_stem)

        if dangling and (i + 1) < n:
            nxt = clean_text(items[i + 1].get("text", ""))
            nxt_stem = strip_q_prefix(nxt)

            if (
                nxt_stem
                and nxt_stem[:1].islower()
                and not LETTERED_OPT_PREFIX_RE.match(nxt_stem)
                and not looks_like_matching_stem(nxt_stem)
                and not NOISE_RE.match(nxt_stem)
                and not STOP_OPTION_RE.match(nxt_stem)
            ):
                combined = clean_text(f"{t_stem} {nxt_stem}")
                if "?" not in combined:
                    combined = combined.rstrip(".") + "?"
                out.append({"text": combined, "is_red": False})
                i += 2
                continue

        out.append(it)
        i += 1

    return out

# ===================================================
# BUILD DESCRIPTION HTML (For learners block)
# ===================================================
def build_description(items):
    collecting = False
    lines = []

    qword_re = re.compile(r"^(what|when|where|which|who|why|how)\b", re.IGNORECASE)

    for it in items:
        t = clean_text(it.get("text",""))
        if not t:
            continue

        if re.search(r"\bFor learners\b|\bFor students\b", t, re.IGNORECASE):
            collecting = True

        # Stop description before the actual question section begins.
        if collecting and re.match(r"^Knowledge\s+Test\b", t, re.IGNORECASE):
            break

        stem_check = strip_q_prefix(t)
        raw_check = clean_text(t)
        looks_like_qmark_stem = (
            stem_check.endswith("?")
            and (
                Q_PREFIX_RE.match(raw_check) is not None
                or NUM_PREFIX_RE.match(raw_check) is not None
                or qword_re.match(stem_check) is not None
            )
        )
        if collecting and (
            QUESTION_CMD_INNER_RE.search(stem_check)
            or looks_like_qmark_stem
            or COMMAND_QUESTION_RE.match(stem_check)
            or re.search(r"\((select|choose)\b", stem_check, re.IGNORECASE)
            or re.search(
                r"^(?:q\s*\d+[\.\)]\s*)?(read\s+the\s+following|complete\s+the)\b",
                stem_check,
                re.IGNORECASE
            )
            or re.search(r"\bdragging\s+and\s+dropping\b|\bdrag\s+and\s+drop\b|\bComplete\s+the\s+table\b", stem_check, re.IGNORECASE)
        ):
            break

        if collecting:
            lines.append(t)

    if not lines:
        return ""

    html_parts = []
    in_list = False

    for ln in lines:
        ln = ln.strip()

        if "•" in ln and not ln.strip().startswith("•"):
            before, *bullets = [p.strip() for p in ln.split("•") if p.strip()]
            if before:
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                html_parts.append(f"<p>{before}</p>")

            if bullets:
                if not in_list:
                    html_parts.append("<ul>")
                    in_list = True
                for b in bullets:
                    html_parts.append(f"<li>{b}</li>")
            continue

        if ln.startswith("•"):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{ln.lstrip('•').strip()}</li>")
            continue

        if in_list:
            html_parts.append("</ul>")
            in_list = False

        html_parts.append(f"<p>{ln}</p>")

    if in_list:
        html_parts.append("</ul>")

    return "\n".join(html_parts)

# ===================================================
# PARSE MCQ QUESTIONS (options + correct in red)
# ===================================================
def parse_mcq_questions(items):
    questions_list = []
    current_q = None
    current_opts = []
    saw_multi_hint = False
    current_start_idx = None
    pending_multi_hint = False

    INSTRUCTION_BLOCK_RE = re.compile(
        r"^(instructions|for\s+learners|for\s+students|for\s+assessors)\b",
        re.IGNORECASE
    )
    META_LINE_RE = re.compile(
        r"^(More than one answer may apply|Select all that apply|Choose all that apply)\b",
        re.IGNORECASE
    )
    META_ANY_RE = re.compile(
        r"\b(More than one answer may apply|Select all that apply|Choose all that apply)\b",
        re.IGNORECASE
    )
    COLON_STEM_RE = re.compile(r":\s*(?:\((?:select|choose)\b.*\))?\s*$", re.IGNORECASE)
    # Keep this narrow to avoid mis-detecting instruction sentences like "When you have completed..."
    # Allow optional "In " prefix (e.g. "In which country...").
    QUESTION_START_RE = re.compile(r"^(?:in\s+)?(which|what|why|how)\b", re.IGNORECASE)
    # Covers stems like "Select the best summary..." / "Choose the most appropriate..."
    SELECT_STEM_RE = re.compile(
        r"^(?:q\s*\d+\.?\s*)?(select|choose|pick)\s+the\s+(best|correct|most\s+appropriate)\b",
        re.IGNORECASE
    )
    READ_STEM_RE = re.compile(r"^(?:q\s*\d+\.?\s*)?read\s+the\s+following\b", re.IGNORECASE)
    COMPLETE_STEM_RE = re.compile(r"^(?:q\s*\d+\.?\s*)?complete\s+the\b", re.IGNORECASE)
    SELECT_HINT_RE = re.compile(r"\((select|choose)\b", re.IGNORECASE)
    FILL_GAP_BLOCK_RE = re.compile(r"\bfill\s+the\s+(gap|blank)\b", re.IGNORECASE)
    CONTAINS_SELECT_SUMMARY_RE = re.compile(
        r"\b(select|choose|pick)\s+the\s+(best|correct|most\s+appropriate)\s+summary\b",
        re.IGNORECASE
    )
    BEST_MATCH_RE = re.compile(
        r"\b(best\s+match|does\s+the\s+following\s+description\s+best\s+match)\b",
        re.IGNORECASE
    )

    def flush():
        nonlocal current_q, current_opts, saw_multi_hint, current_start_idx
        if not current_q:
            return

        opts = [o for o in current_opts if not NOISE_RE.match(o["text"])]
        opts = [o for o in opts if not OPTION_NOISE_RE.match(o["text"])]
        option_texts = [o["text"] for o in opts]
        correct = [i for i, o in enumerate(opts) if o["is_red"]]

        qtext = strip_q_prefix(current_q.strip())
        qlower = qtext.lower()
        multi = (
            saw_multi_hint
            or bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower))
            or ("apply" in qlower)
            or (len(correct) > 1)
        )

        def is_strong_stem(txt: str) -> bool:
            t = (txt or "").strip()
            return bool(
                t.endswith("?")
                or QUESTION_CMD_INNER_RE.search(t)
                or SELECT_HINT_RE.search(t)
                or META_ANY_RE.search(t)
                or COLON_STEM_RE.search(t)
                or SELECT_STEM_RE.match(t)
                or QUESTION_START_RE.match(t)
            )

        # If options weren't extracted (often due to images/shapes), keep a placeholder only
        # when the stem is very likely a real question (avoid turning rubric text into MCQs).
        if len(option_texts) < 2:
            if not is_strong_stem(qtext):
                current_q = None
                current_opts = []
                saw_multi_hint = False
                current_start_idx = None
                return

            option_texts = [
                "⚠ Option text not extracted (likely image/shape). Please replace this option.",
                "⚠ Option text not extracted (likely image/shape). Please replace this option.",
            ]
            correct = []

        questions_list.append({
            "question": qtext,
            "options": option_texts,
            "correct": correct,
            "multi": multi,
            "kind": "mcq",
            "_order": (current_start_idx if current_start_idx is not None else 10**9),
        })

        current_q = None
        current_opts = []
        saw_multi_hint = False
        current_start_idx = None

    def parse_fill_gap_line(line: str) -> tuple[str, list[str]] | None:
        if line.count("/") < 2:
            return None

        parts = [p.strip() for p in re.split(r"\s*/\s*", line) if p.strip()]
        if len(parts) < 3:
            return None

        opt0 = parts[0].split()[-1]
        opt_last = parts[-1].split()[0]

        if not opt0 or not opt_last:
            return None

        prefix = parts[0][: -len(opt0)].rstrip()
        suffix = parts[-1][len(opt_last):].lstrip()

        options = [opt0] + parts[1:-1] + [opt_last]
        options = [clean_text(o) for o in options if clean_text(o)]

        qtext = clean_text(f"{prefix} ____ {suffix}".strip())
        if len(qtext) < 10 or len(options) < 3:
            return None

        return qtext, options

    def has_plausible_options(start_idx: int) -> bool:
        """
        Heuristic guard so we don't treat random question-mark sentences as MCQ stems.
        """
        n = len(items)
        count = 0
        for j in range(start_idx, min(n, start_idx + 25)):
            raw = clean_text(items[j].get("text", ""))
            if not raw or NOISE_RE.match(raw):
                continue
            t = strip_q_prefix(raw)

            if ESSAY_GUIDE_RE.match(t) or RUBRIC_START_RE.match(t):
                return False
            if looks_like_matching_stem(t):
                return False
            if SELECT_STEM_RE.match(t) or QUESTION_CMD_INNER_RE.search(t) or BEST_MATCH_RE.search(t) or CONTAINS_SELECT_SUMMARY_RE.search(t):
                break
            if SELECT_HINT_RE.search(t):
                break
            if t.endswith("?") and len(t) >= 10:
                break

            if len(t) <= 200 and not t.endswith("?"):
                count += 1
                if count >= 2:
                    return True
        return False

    for idx, it in enumerate(items):
        line = clean_text(it.get("text",""))
        if not line:
            continue
        if NOISE_RE.match(line):
            continue
        if OPTION_NOISE_RE.match(line):
            continue

        # If we hit an instructions block mid-parse, don't treat it as an MCQ or options.
        if INSTRUCTION_BLOCK_RE.match(strip_q_prefix(line)):
            flush()
            current_q = None
            current_opts = []
            saw_multi_hint = False
            current_start_idx = None
            pending_multi_hint = False
            continue
        if ESSAY_GUIDE_RE.match(line):
            current_q = None
            current_opts = []
            continue
        if current_q and STOP_OPTION_RE.match(line):
            flush()
            current_q = None
            current_opts = []
            continue

        t_stem = strip_q_prefix(line)

        # Meta line sometimes appears on its own just before the question/options (or at the end).
        # Don't turn it into a question; carry it forward as a hint for the next stem.
        if current_q is None and META_LINE_RE.match(t_stem):
            pending_multi_hint = True
            continue

        # Never parse matching instructions as MCQ.
        if looks_like_matching_stem(t_stem):
            flush()
            current_q = None
            current_opts = []
            saw_multi_hint = False
            current_start_idx = None
            continue

        # Stems that end with ":" (common in assessor guides), including "(select all that apply)".
        if (
            (current_q is None or len(current_opts) >= 2)
            and COLON_STEM_RE.search(t_stem)
            and len(t_stem) >= 12
            and not looks_like_matching_stem(t_stem)
            and not COMMAND_QUESTION_RE.match(t_stem)
            and not STOP_OPTION_RE.match(line)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = bool(META_ANY_RE.search(t_stem) or pending_multi_hint)
            current_start_idx = idx
            pending_multi_hint = False
            continue

        # Stems that include "select all that apply" even if the options are in a non-extractable shape.
        if (
            (current_q is None or len(current_opts) >= 2)
            and META_ANY_RE.search(t_stem)
            and not META_LINE_RE.match(t_stem)
            and len(t_stem) >= 12
            and not looks_like_matching_stem(t_stem)
            and not COMMAND_QUESTION_RE.match(t_stem)
            and not STOP_OPTION_RE.match(line)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = True
            current_start_idx = idx
            pending_multi_hint = False
            continue

        # Some assessor guides include a "Read the following... fill the gap..." instruction line
        # followed by multiple slash-option statements. Don't treat the instruction as an MCQ;
        # flush any current question and let the slash-lines become their own questions.
        if FILL_GAP_BLOCK_RE.search(t_stem):
            flush()
            current_q = None
            current_opts = []
            saw_multi_hint = False
            current_start_idx = None
            continue

        if (current_q is None) and ("/" in t_stem) and has_plausible_options(idx + 1):
            parsed = parse_fill_gap_line(t_stem)
            if parsed:
                qtext, opts = parsed
                questions_list.append({
                    "question": qtext,
                    "options": opts,
                    "correct": [],
                    "multi": False,
                    "kind": "mcq",
                    "_order": idx,
                })
                continue

        if (
            SELECT_HINT_RE.search(t_stem)
            and (current_q is None or len(current_opts) >= 2)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if (
            (READ_STEM_RE.match(t_stem) or COMPLETE_STEM_RE.match(t_stem))
            and ("select" in t_stem.lower() or "most appropriate" in t_stem.lower() or "complete" in t_stem.lower())
            and not looks_like_matching_stem(t_stem)
            and (current_q is None or len(current_opts) >= 2)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        # New question stem (Select/Choose the best/correct/most appropriate ...)
        if (
            SELECT_STEM_RE.match(line)
            and not looks_like_matching_stem(line)
            and not COMMAND_QUESTION_RE.match(strip_q_prefix(line))
            and not META_LINE_RE.match(line)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        m = QUESTION_CMD_INNER_RE.search(line)
        if m:
            flush()
            start = m.start()
            stem = line[:start].strip()
            cmd_plus = line[start:].strip()
            q_line = f"{stem} {cmd_plus}".strip() if stem else cmd_plus
            current_q = strip_q_prefix(q_line)
            current_opts = []
            current_start_idx = idx
            saw_multi_hint = pending_multi_hint
            pending_multi_hint = False
            continue

        if current_q and not current_opts:
            if (
                current_q and not current_q.strip().endswith("?")
                and line[:1].islower()
                and not QUESTION_CMD_INNER_RE.search(line)
                and not looks_like_matching_stem(line)
                and not COMMAND_QUESTION_RE.match(strip_q_prefix(line))
                and not STOP_OPTION_RE.match(line)
                and not META_LINE_RE.match(line)
            ):
                current_q = (current_q + " " + line).strip()
                continue

        if (
            QUESTION_START_RE.match(t_stem)
            and len(t_stem) >= 12
            and not looks_like_matching_stem(t_stem)
            and not COMMAND_QUESTION_RE.match(t_stem)
            and not STOP_OPTION_RE.match(line)
            and not META_LINE_RE.match(line)
            and (
                "?" in t_stem
                or BEST_MATCH_RE.search(t_stem)
                or CONTAINS_SELECT_SUMMARY_RE.search(t_stem)
                or re.search(r"\((select|choose)\b", t_stem, re.IGNORECASE)
            )
            and (current_q is None or len(current_opts) >= 2)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        # Stems that don't start with question words (e.g. "In regard to ..., select the best summary ...")
        if (
            (CONTAINS_SELECT_SUMMARY_RE.search(t_stem) or BEST_MATCH_RE.search(t_stem))
            and (current_q is None or len(current_opts) >= 2)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if current_q and META_LINE_RE.match(line):
            saw_multi_hint = True
            continue

        t = strip_q_prefix(line)
        if (
            t.endswith("?")
            and len(t) >= 10
            and not COMMAND_QUESTION_RE.match(t)
            and not looks_like_matching_stem(t)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t
            current_opts = []
            current_start_idx = idx
            continue

        if current_q:
            current_opts.append({"text": line, "is_red": it.get("is_red", False)})

    flush()
    return [q for q in questions_list if len(q.get("options") or []) >= 2 and len(q.get("question") or "") >= 10]

# ===================================================
# PARSE ESSAY QUESTIONS
# ===================================================
def parse_essay_questions(items):
    questions = []
    n = len(items)

    i = 0
    while i < n:
        raw = clean_text(items[i].get("text",""))
        if not raw or NOISE_RE.match(raw):
            i += 1
            continue

        line = strip_q_prefix(raw)

        if COMMAND_QUESTION_RE.match(line):
            j = i + 1
            next_line = ""
            while j < n:
                nxt = clean_text(items[j].get("text",""))
                if nxt and not NOISE_RE.match(nxt):
                    next_line = nxt
                    break
                j += 1

            if RUBRIC_START_RE.match(next_line):
                questions.append({
                    "question": line,
                    "options": [],
                    "correct": [],
                    "multi": False,
                    "kind": "essay",
                    "_order": i,
                })
                i = j + 1
                continue

        i += 1

    return [q for q in questions if len((q.get("question") or "").strip()) >= 10]


# ===================================================
# MATCHING PARSER (FULL)  ✅ paste this BEFORE you call it
# ===================================================

# ===================================================
# MATCHING PARSER (RECUSRIVE + DEDUPE + BULLET SAFE)
# - Finds matching tables even when nested inside layout tables
# - Dedupes by fingerprint so nested traversal doesn't multiply tables
# - Keeps multi-line / bullet list cells (joins with "; ")
# - Skips instruction/rubric tables more safely
# - If stem not found, creates a sensible fallback from table headers
# ===================================================

MATCHING_STEM_RE = re.compile(
    r"\b("
    r"complete\s+the\s+table|"
    r"drag(?:ging)?\s+and\s+drop(?:ping)?|"
    r"drag\s+and\s+drop|"
    r"match\s+each|"
    r"match\s+the\s+following|"
    r"match\s+.*\s+to\s+the\s+correct|"
    r"select\s+one.*for\s+each"
    r")\b",
    re.IGNORECASE
)

INSTRUCTION_TABLE_NOISE_RE = re.compile(
    r"\b("
    r"range\s+and\s+conditions?|decision-?making\s+rules?|"
    r"rubric|pre-?approved\s+reasonable\s+adjustments?|"
    r"for\s+learners?|for\s+assessors?|instructions?|"
    r"learner\s+feedback|assessment\s+outcome|assessor\s+signature|assessor\s+name|final\s+comments|date|"
    r"evidence|required|criteria|competent|nyc|submission|marking"
    r")\b",
    re.IGNORECASE
)

INSTRUCTION_ROW_LABEL_RE = re.compile(
    r"^(instructions|for\s+learners:?|for\s+assessors:?|range\s+and\s+conditions|decision-?making\s+rules|"
    r"pre-?approved\s+reasonable\s+adjustments|rubric|learner\s+feedback|assessment\s+outcome|"
    r"assessor\s+signature|assessor\s+name|final\s+comments|date)\b",
    re.IGNORECASE
)

def looks_like_matching_stem(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors")):
        return False
    if COMMAND_QUESTION_RE.match(t2):  # avoid essay stems stealing matching
        return False
    if "which of the following" in low:  # avoid mcq stems
        return False
    return bool(MATCHING_STEM_RE.search(t2))

def join_lines(lines: list[str]) -> str:
    parts = [clean_text(x) for x in (lines or []) if clean_text(x)]
    return "; ".join(parts).strip()

OPTION_LABEL_PREFIX_RE = re.compile(r"^\s*(?:[\(\[]?[a-zA-Z0-9]{1,3}[\)\].:-])\s+")
def parse_pasted_options(blob: str) -> list[str]:
    lines = []
    for raw in (blob or "").splitlines():
        t = clean_text(raw)
        if not t:
            continue
        t = OPTION_LABEL_PREFIX_RE.sub("", t).strip()
        if t:
            lines.append(t)

    # de-dupe keep order
    out, seen = [], set()
    for x in lines:
        k = normalize_key(x)
        if k in seen:
            continue
        seen.add(k)
        out.append(x)
    return out

def cell_lines(cell) -> list[str]:
    """
    Bullet-safe cell extraction:
    - Each paragraph becomes a line (works even when Word bullets aren't '•')
    - If a paragraph contains '•' inline, we split it too.
    """
    lines = []
    for p in cell.paragraphs:
        t, _ = paragraph_text_and_is_red(p)
        t = clean_text(t)
        if not t:
            continue

        # If Word exported bullet symbol in same paragraph
        if "•" in t:
            parts = [x.strip() for x in t.split("•") if x.strip()]
            lines.extend(parts)
        else:
            lines.append(t)

    # de-dupe keep order
    out, seen = [], set()
    for x in lines:
        x = clean_text(x)
        if not x:
            continue
        k = normalize_key(x)
        if k in seen:
            continue
        seen.add(k)
        out.append(x)
    return out

def table_to_grid(table: Table) -> list[list[list[str]]]:
    return [[cell_lines(c) for c in row.cells] for row in table.rows]

def table_fingerprint(grid) -> str:
    flat = []
    for row in grid:
        for cell in row:
            flat.append("|".join(cell))
    blob = "||".join([normalize_key(x) for x in flat if x])
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()

def is_instruction_table(grid) -> bool:
    """
    Skip obvious policy/rubric tables.
    BUT don't be too aggressive (so we don't kill real matching).
    """
    texts = []
    for row in grid:
        for cell in row:
            if cell:
                texts.append(" ".join(cell))
    blob = " ".join(texts).strip()
    if not blob:
        return True

    # If the first row looks like a policy header table
    first_row = " ".join([join_lines(c) for c in (grid[0] if grid else [])]).lower()
    if "range and conditions" in first_row or "decision-making rules" in first_row or "pre-approved" in first_row:
        return True

    # If the first column contains multiple rubric/policy labels, it's not a matching question
    label_hits = 0
    for row in grid[: min(len(grid), 12)]:
        if not row:
            continue
        left = join_lines(row[0]).strip()
        if left and INSTRUCTION_ROW_LABEL_RE.search(left):
            label_hits += 1
            if label_hits >= 2:
                return True

    hits = sum(1 for t in texts if INSTRUCTION_TABLE_NOISE_RE.search(t))
    ratio = hits / max(1, len(texts))

    # only skip if *many* cells look like policy text
    return ratio >= 0.40

def guess_header_skip(grid) -> int:
    if not grid or not grid[0]:
        return 0

    # Join first row text
    row0 = " ".join([join_lines(c) for c in grid[0] if c]).strip().lower()

    # Common header words (add anything you see in your files)
    header_words = [
        "definition", "term", "meaning", "word", "concept",
        "numbers", "number", "example", "type", "classification",
        "left", "right"
    ]

    # If first row looks like headers, skip it
    if any(w in row0 for w in header_words):
        return 1

    # Also skip if it's mostly very short labels (typical header row)
    nonempty = [join_lines(c) for c in grid[0] if join_lines(c)]
    if nonempty and sum(len(x) <= 20 for x in nonempty) / len(nonempty) >= 0.8:
        return 1

    return 0


def pair_is_valid(left: str, right: str) -> bool:
    if not left or not right:
        return False
    if normalize_key(left) == normalize_key(right):
        return False
    # allow longer text; instruction tables are filtered elsewhere
    if len(left) > 450 or len(right) > 900:
        return False
    return True

def score_columns(grid, a: int, b: int) -> int:
    start = guess_header_skip(grid)
    score = 0
    for r in range(start, len(grid)):
        if a >= len(grid[r]) or b >= len(grid[r]):
            continue
        left = join_lines(grid[r][a])
        right = join_lines(grid[r][b])
        if pair_is_valid(left, right):
            score += 1
    return score

def pick_best_columns(grid) -> tuple[int, int] | None:
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    if max_cols < 2:
        return None

    best, best_score = None, 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = score_columns(grid, a, b)
            if sc > best_score:
                best_score = sc
                best = (a, b)

    # require at least 2 usable pairs
    if best is None or best_score < 2:
        return None
    return best


def extract_pairs(grid, left_col: int, right_col: int, start_row: int = 0) -> list[dict]:
    pairs = []

    for r in range(start_row, len(grid)):
        if left_col >= len(grid[r]) or right_col >= len(grid[r]):
            continue

        left = join_lines(grid[r][left_col])
        right = join_lines(grid[r][right_col])

        left = re.sub(r"^\(?[a-z]\)\s*", "", left, flags=re.IGNORECASE).strip()
        left = re.sub(r"^[a-z]\.\s*", "", left, flags=re.IGNORECASE).strip()

        if not pair_is_valid(left, right):
            continue

        pairs.append({"left": left, "right": right})

    seen, out = set(), []
    for p in pairs:
        k = normalize_key(p["left"]) + "=>" + normalize_key(p["right"])
        if k in seen:
            continue
        seen.add(k)
        out.append(p)
    return out


def iter_elements_recursive(container):
    """
    Yield paragraphs + tables in document order, recursively inside tables.
    SAFE: no Document/DocxDocument isinstance checks.
    """
    # If it has `.element`, it's a Document-like object
    if hasattr(container, "element"):
        parent_elm = container.element.body
        parent_obj = container
    else:
        # table cell
        parent_elm = container._tc
        parent_obj = container

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", Paragraph(child, parent_obj))
        elif child.tag.endswith("}tbl"):
            tbl = Table(child, parent_obj)
            yield ("tbl", tbl)
            for row in tbl.rows:
                for cell in row.cells:
                    yield from iter_elements_recursive(cell)


def parse_matching_questions_doc_order(docx_path: str, item_index: dict[str, int] | None = None) -> list[dict]:
    doc = Document(docx_path)
    out = []
    recent_paras: list[str] = []
    recent_orders: list[int] = []
    seen_tables = set()
    MAX_LOOKBACK = 50
    seq = 0

    def choose_stem() -> tuple[str | None, int | None]:
        paras = recent_paras[-MAX_LOOKBACK:]
        orders = recent_orders[-MAX_LOOKBACK:] if recent_orders else [None] * len(paras)
        for t, o in reversed(list(zip(paras, orders))):
            if looks_like_matching_stem(t):
                return strip_q_prefix(clean_text(t)), o
        return None, None

    for kind, el in iter_elements_recursive(doc):
        seq += 1
        if kind == "p":
            t, _ = paragraph_text_and_is_red(el)
            t = clean_text(t)
            if t:
                recent_paras.append(t)
                if item_index:
                    recent_orders.append(item_index.get(normalize_key(t), seq))
                else:
                    recent_orders.append(seq)
                if len(recent_paras) > 400:
                    recent_paras = recent_paras[-400:]
                    recent_orders = recent_orders[-400:]
            continue

        # table
        grid = table_to_grid(el)
        header_skip = guess_header_skip_by_row_color(el)  # ✅ uses row shading

        tfp = table_fingerprint(grid)
        if tfp in seen_tables:
            continue
        seen_tables.add(tfp)

        if is_instruction_table(grid):
            continue

        # ✅ If one column is shaded (Term column), force it to be LEFT
        term_col = pick_term_column_by_fill(el)

        if term_col is not None:
    # choose RIGHT column by best score against the forced LEFT
            max_cols = max(len(r) for r in grid)
            best_right = None
            best_score = 0
            for b in range(max_cols):
                if b == term_col:
                    continue
                sc = score_columns(grid, term_col, b)
                if sc > best_score:
                    best_score = sc
                    best_right = b

            if best_right is None or best_score < 2:
                continue

            left_col, right_col = term_col, best_right

        else:
            cols = pick_best_columns(grid)
            if not cols:
                continue
            left_col, right_col = cols

        pairs = extract_pairs(grid, left_col, right_col, start_row=header_skip)

        if len(pairs) < 2:
            continue

        # prefer nearby stem, but DO NOT require it
        stem, stem_order = choose_stem()
        if not stem:
            # fallback stem from header row if present
            header = grid[0] if grid else []
            hL = (header[left_col][0] if header and left_col < len(header) and header[left_col] else "Left")
            hR = (header[right_col][0] if header and right_col < len(header) and header[right_col] else "Right")
            stem = f"Match each '{hL}' to the correct '{hR}'."
        if stem_order is None:
            stem_order = seq

        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False,
            "_order": stem_order,
        })

        # reset so the next table doesn't steal the same stem
        recent_paras = []
        recent_orders = []

    return out


def cell_fill_hex(cell) -> str | None:
    """
    Returns background fill like 'D9D9D9' or None if no shading.
    """
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill:
        return None
    fill = fill.strip().upper()
    if fill in ("AUTO", "FFFFFF"):  # treat white as "no shading"
        return None
    return fill

def row_fill_signature(row_cells) -> tuple:
    # tuple of fills per cell so we can compare rows
    return tuple(cell_fill_hex(c) for c in row_cells)

def guess_header_skip_by_row_color(table: Table) -> int:
    """
    If row 0 is shaded (grey) and most body rows are not, skip it.
    """
    if not table.rows:
        return 0

    row0 = table.rows[0]
    sig0 = row_fill_signature(row0.cells)

    # If header row has ANY fill, and the next few rows mostly don't → header
    if any(sig0):
        sample = table.rows[1: min(6, len(table.rows))]
        non_header_like = 0
        for r in sample:
            sig = row_fill_signature(r.cells)
            if not any(sig):  # no fill
                non_header_like += 1
        if non_header_like >= max(1, len(sample) - 1):
            return 1

    return 0

def column_fill_stats(table: Table):
    """
    For each column index, count how many cells have non-empty fill.
    Returns: dict[col_index] = count
    """
    stats = {}
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            f = cell_fill_hex(cell)   # ✅ correct function name
            if f:
                stats[ci] = stats.get(ci, 0) + 1
    return stats



def pick_term_column_by_fill(table: Table) -> int | None:
    """
    If one column is clearly shaded more than others, treat it as the Term column.
    """
    stats = column_fill_stats(table)
    if not stats:
        return None

    # pick the column with most shaded cells
    best_col = max(stats, key=lambda c: stats[c])
    best = stats[best_col]

    # require it to be "clearly" more shaded than next best
    sorted_counts = sorted(stats.values(), reverse=True)
    second = sorted_counts[1] if len(sorted_counts) > 1 else 0

    # tune thresholds if needed
    if best >= 2 and best >= second + 2:
        return best_col
    return None

# ===================================================
# CANVAS API HELPERS
# ===================================================
def canvas_headers(canvas_token: str):
    return {"Authorization": f"Bearer {canvas_token}"}

def canvas_whoami(canvas_base_url: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/users/self"
    r = requests.get(url, headers=canvas_headers(canvas_token), timeout=30)
    if r.status_code == 401:
        return None
    r.raise_for_status()
    return r.json()

def list_courses(canvas_base_url: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/courses"
    out = []
    page = 1
    while True:
        r = requests.get(
            url,
            headers=canvas_headers(canvas_token),
            params={"per_page": 100, "page": page},
            timeout=60
        )
        r.raise_for_status()
        batch = r.json()
        if not batch:
            break
        out.extend(batch)
        page += 1
    return out

def get_existing_quiz_titles(canvas_base_url: str, course_id: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes"
    titles = set()
    page = 1
    while True:
        r = requests.get(
            url,
            headers=canvas_headers(canvas_token),
            params={"page": page, "per_page": 100},
            timeout=60
        )
        r.raise_for_status()
        data = r.json()
        if not data:
            break
        for q in data:
            titles.add((q.get("title") or "").strip())
        page += 1
    return titles

def generate_unique_title(base_title, existing_titles):
    if base_title not in existing_titles:
        return base_title
    i = 1
    while True:
        candidate = f"{base_title} ({i})"
        if candidate not in existing_titles:
            return candidate
        i += 1

def create_canvas_quiz(canvas_base_url: str, course_id: str, canvas_token: str,
                       title: str, description_html: str = "", settings: dict | None = None) -> int:
    settings = settings or {}

    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes"
    quiz_obj = {
        "title": title,
        "description": description_html,
        "published": False,
        "quiz_type": "assignment",
        "shuffle_answers": bool(settings.get("shuffle_answers", False)),
        "one_question_at_a_time": bool(settings.get("one_question_at_a_time", False)),
        "show_correct_answers": bool(settings.get("show_correct_answers", False)),
        "scoring_policy": settings.get("scoring_policy", "keep_highest"),
    }

    tl = int(settings.get("time_limit", 0) or 0)
    if tl > 0:
        quiz_obj["time_limit"] = tl

    allow_multi = bool(settings.get("allow_multiple_attempts", False))
    if allow_multi:
        aa = int(settings.get("allowed_attempts", 2) or 2)
        quiz_obj["allowed_attempts"] = max(2, aa)

    if bool(settings.get("access_code_enabled", False)) and (settings.get("access_code") or "").strip():
        quiz_obj["access_code"] = settings["access_code"].strip()

    for k in ["due_at", "unlock_at", "lock_at"]:
        v = (settings.get(k) or "").strip()
        if v:
            quiz_obj[k] = v

    payload = {"quiz": quiz_obj}

    r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    if r.status_code == 401:
        raise RuntimeError("401 Unauthorized — token invalid/expired.")
    if r.status_code == 403:
        raise RuntimeError("403 Forbidden — missing permission in this course.")
    r.raise_for_status()
    return r.json()["id"]

def publish_quiz(canvas_base_url: str, course_id: str, canvas_token: str, quiz_id: int):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes/{quiz_id}"
    payload = {"quiz": {"published": True}}
    r = requests.put(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    r.raise_for_status()

def add_question_to_quiz(canvas_base_url: str, course_id: str, canvas_token: str, quiz_id: int, q: dict):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes/{quiz_id}/questions"
    qtext = strip_q_prefix((q.get("question") or "").strip())
    kind = (q.get("kind") or "").lower()

    # MATCHING
    if kind == "matching":
        pairs = q.get("pairs") or []
        answers = []
        for p in pairs:
            left = (p.get("left") or "").strip()
            right = (p.get("right") or "").strip()
            if left and right:
                answers.append({
                    "answer_match_left": left,
                    "answer_match_right": right,
                    "answer_weight": 100
                })

        payload = {
            "question": {
                "question_name": qtext[:100] if qtext else "Matching",
                "question_text": qtext,
                "question_type": "matching_question",
                "points_possible": 1,
                "answers": answers
            }
        }
        r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
        if r.status_code >= 400:
            raise RuntimeError(f"Canvas error {r.status_code}: {r.text[:600]}")
        r.raise_for_status()
        return

    # ESSAY
    opts = [o.strip() for o in (q.get("options") or []) if o and o.strip()]
    correct = q.get("correct", []) or []
    if kind == "essay" or len(opts) < 2:
        payload = {
            "question": {
                "question_name": qtext[:100] if qtext else "Question",
                "question_text": qtext or " ",
                "question_type": "essay_question",
                "points_possible": 1
            }
        }
        r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
        r.raise_for_status()
        return

    # MCQ
    qlower = (qtext or "").lower()
    multi = bool(q.get("multi")) or (len(correct) > 1) or bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower)) or ("apply" in qlower)
    qtype = "multiple_answers_question" if multi else "multiple_choice_question"

    answers = [{"answer_text": opt, "answer_weight": 100 if idx in correct else 0} for idx, opt in enumerate(opts)]

    payload = {
        "question": {
            "question_name": (qtext[:100] if qtext else "Question"),
            "question_text": qtext,
            "question_type": qtype,
            "points_possible": 1,
            "answers": answers
        }
    }

    r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    if r.status_code >= 400:
        raise RuntimeError(f"Canvas error {r.status_code}: {r.text[:600]}")
    r.raise_for_status()

# ===================================================
# SIDEBAR: LOGIN + COURSE PICKER
# ===================================================
with st.sidebar:
    st.header("🔐 Canvas Login (Token)")

    st.session_state.canvas_base_url = st.text_input(
        "Canvas Base URL",
        value=st.session_state.canvas_base_url,
        help="Example: https://learningvault.test.instructure.com/api/v1"
    ).strip()

    st.session_state.canvas_token = st.text_input(
        "Canvas Access Token",
        value=st.session_state.canvas_token,
        type="password",
        help="Canvas API uses access tokens."
    )

    c_login, c_logout = st.columns(2)

    if c_login.button("Login"):
        try:
            me = canvas_whoami(st.session_state.canvas_base_url, st.session_state.canvas_token)
            if me:
                st.session_state.logged_in = True
                st.session_state.me = me
                st.session_state.courses_cache = None
                st.success(f"Logged in as: {me.get('name', 'Unknown')}")
            else:
                st.session_state.logged_in = False
                st.session_state.me = None
                st.error("Login failed: token invalid/expired.")
        except Exception as e:
            st.session_state.logged_in = False
            st.session_state.me = None
            st.error(f"Login failed: {e}")

    if c_logout.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.me = None
        st.session_state.selected_course_id = None
        st.session_state.courses_cache = None
        st.session_state.questions = []
        st.session_state.parsed_ok = False
        st.info("Logged out.")

    st.divider()

    if st.session_state.logged_in:
        st.subheader("✅ Course")
        try:
            if st.session_state.courses_cache is None:
                st.session_state.courses_cache = list_courses(
                    st.session_state.canvas_base_url,
                    st.session_state.canvas_token
                )

            courses = st.session_state.courses_cache or []
            if not courses:
                st.warning("No courses visible to this token.")
            else:
                label_to_id = {}
                labels = []
                for c in courses:
                    cid = c.get("id")
                    name = (c.get("name") or c.get("course_code") or f"Course {cid}").strip()
                    label = f"{name} (ID: {cid})"
                    labels.append(label)
                    label_to_id[label] = str(cid)

                default_index = 0
                if st.session_state.selected_course_id:
                    for i, lb in enumerate(labels):
                        if label_to_id[lb] == st.session_state.selected_course_id:
                            default_index = i
                            break

                chosen = st.selectbox("Select course", labels, index=default_index)
                st.session_state.selected_course_id = label_to_id[chosen]

        except Exception as e:
            st.error(f"Failed to load courses: {e}")
    else:
        st.info("Login first to select course and upload.")

# Block main area until logged in + course chosen
if not st.session_state.logged_in:
    st.warning("Please login in the sidebar first.")
    st.stop()

if not st.session_state.selected_course_id:
    st.warning("Please select a course in the sidebar.")
    st.stop()

course_id = st.session_state.selected_course_id
canvas_base_url = st.session_state.canvas_base_url
canvas_token = st.session_state.canvas_token

# ===================================================
# MAIN UI: UPLOAD + PARSE
# ===================================================
st.subheader("1) Upload DOCX and Parse")

uploaded = st.file_uploader("📤 Upload DOCX assessment file", type=["docx"])
colA, colB = st.columns([1, 1])
parse_btn = colA.button("🧠 Parse DOCX", type="primary", disabled=(uploaded is None))
reset_btn = colB.button("♻️ Reset parsed data")

log_box = st.empty()

if reset_btn:
    st.session_state.docx_filename = None
    st.session_state.description_html = ""
    st.session_state.questions = []
    st.session_state.parsed_ok = False
    st.success("Reset complete.")

if parse_btn:
    if not uploaded:
        st.warning("Upload a DOCX first.")
        st.stop()

    # ✅ clear previous results FIRST (prevents “previous answers carry over”)
    st.session_state.questions = []
    st.session_state.parsed_ok = False
    st.session_state.description_html = ""
    st.session_state.docx_filename = None
        # ✅ bump parse run id so Streamlit forgets old widget values
    st.session_state.parse_run_id += 1


    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(uploaded.read())
        docx_path = tmp.name

    buf = io.StringIO()
    try:
        with contextlib.redirect_stdout(buf):
            items = extract_items_with_red(docx_path)
            items = merge_dangling_question_lines(items)
            description_html = build_description(items)

            item_index: dict[str, int] = {}
            for i, it in enumerate(items):
                key = normalize_key(clean_text(it.get("text", "")))
                if key and key not in item_index:
                    item_index[key] = i

            matching = parse_matching_questions_doc_order(docx_path, item_index=item_index)
            mcq = parse_mcq_questions(items)
            mcq = repair_placeholder_mcq_options(mcq, items, item_index)
            essay = parse_essay_questions(items)

            questions = matching + mcq + essay
            questions.sort(key=lambda q: int(q.get("_order", 10**9)))
            questions = dedupe_questions(questions)
            questions = collapse_duplicate_mcq(questions)

            print("DEBUG: items extracted:", len(items))
            print("DEBUG: matching:", len(matching), "mcq:", len(mcq), "essay:", len(essay))
            print("Parsed questions:", len(questions))

            debug_phrases = [
                "Description of food item",
                "Use By or Best Before date",
                "Balance",
                "Colour",
                "Contrast",
                "Assemble and/or set up required equipment and utensils",
                "Obtain, weigh or measure ingredients",
            ]
            for ph in debug_phrases:
                in_items = any(ph.lower() in (it.get("text", "").lower()) for it in items)
                if in_items:
                    continue
                in_xml = docx_xml_contains(docx_path, ph)
                if in_xml:
                    print(f"DEBUG: phrase present in DOCX XML but not extracted: {ph!r}")
                else:
                    print(f"DEBUG: phrase not in DOCX XML (likely an image): {ph!r}")

            missing_phrase_debug = []
            for ph in debug_phrases:
                in_items = any(ph.lower() in (it.get("text", "").lower()) for it in items)
                if in_items:
                    continue
                missing_phrase_debug.append({
                    "phrase": ph,
                    "in_docx_xml": bool(docx_xml_contains(docx_path, ph)),
                })

            if not questions:
                raise RuntimeError("No questions detected. Check: headings, question wording, or doc format.")

        st.session_state.docx_filename = uploaded.name
        st.session_state.description_html = description_html
        st.session_state.questions = questions
        st.session_state.missing_phrase_debug = missing_phrase_debug
        st.session_state.parsed_ok = True
        st.session_state.debug_docx_copy = None

        # Keep a copy of the uploaded DOCX for local debugging when we detect extraction problems.
        # (e.g. options rendered inside shapes/text boxes or other non-paragraph structures)
        has_placeholder_opts = any(
            any((o or "").startswith("⚠ Option text not extracted") for o in (q.get("options") or []))
            for q in questions
            if (q.get("kind") or "").lower() == "mcq"
        )
        if st.session_state.missing_phrase_debug or has_placeholder_opts:
            saver = globals().get("save_docx_debug_copy")
            if callable(saver):
                st.session_state.debug_docx_copy = saver(docx_path, uploaded.name)

        log_box.code(buf.getvalue())
        st.success(f"✅ Parsed {len(questions)} questions.")
        img_like = [x["phrase"] for x in st.session_state.missing_phrase_debug if not x.get("in_docx_xml")]
        xml_like = [x["phrase"] for x in st.session_state.missing_phrase_debug if x.get("in_docx_xml")]
        st.caption(f"Missing-phrase debug: {len(st.session_state.missing_phrase_debug)} missing")
        if img_like:
            st.warning("Some expected labels look like images (not extractable without OCR): " + ", ".join(img_like))
        if xml_like:
            st.warning("Some expected labels are in the DOCX but still not extracted (needs parser update): " + ", ".join(xml_like))
        if st.session_state.debug_docx_copy:
            st.caption(f"Saved debug copy: {st.session_state.debug_docx_copy}")
            try:
                with open(st.session_state.debug_docx_copy, "rb") as f:
                    st.download_button(
                        "Download debug DOCX copy",
                        data=f.read(),
                        file_name=os.path.basename(st.session_state.debug_docx_copy),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            except Exception:
                pass

    except Exception as e:
        log_box.code(buf.getvalue())
        st.error(f"❌ Parse error: {e}")
        with st.expander("Show traceback"):
            st.code(traceback.format_exc())

    finally:
        try:
            os.remove(docx_path)
        except Exception:
            pass

# ===================================================
# STOP if nothing parsed yet
# ===================================================
if not st.session_state.parsed_ok:
    st.info("Parse a DOCX to continue (you can still select course and login first).")
    st.stop()

questions = st.session_state.questions

st.divider()

# ===================================================
# DETAILS
# ===================================================
st.subheader("2) Details (Canvas Quiz Settings)")

default_title = os.path.splitext(st.session_state.docx_filename or "Quiz")[0]
quiz_title = st.text_input("Quiz Title *", value=default_title)

quiz_instructions = st.text_area(
    "Quiz Instructions (HTML allowed)",
    value=st.session_state.description_html or "",
    height=160
)

d = st.session_state.details

c1, c2, c3 = st.columns(3)
d["shuffle_answers"] = c1.checkbox("Shuffle Answers", value=bool(d.get("shuffle_answers", True)))
d["one_question_at_a_time"] = c2.checkbox("Show one question at a time", value=bool(d.get("one_question_at_a_time", False)))
d["show_correct_answers"] = c3.checkbox("Let Students See The Correct Answers", value=bool(d.get("show_correct_answers", False)))

c4, c5, c6 = st.columns(3)
d["time_limit"] = c4.number_input("Time Limit (minutes, 0 = none)", min_value=0, value=int(d.get("time_limit", 0)), step=5)
d["allow_multiple_attempts"] = c5.checkbox("Allow Multiple Attempts", value=bool(d.get("allow_multiple_attempts", False)))

if d["allow_multiple_attempts"]:
    current_attempts = int(d.get("allowed_attempts", 2) or 2)
    if current_attempts < 2:
        current_attempts = 2
        d["allowed_attempts"] = 2
    d["allowed_attempts"] = c5.number_input("Allowed Attempts", min_value=2, value=current_attempts, step=1)
else:
    d["allowed_attempts"] = 1

d["scoring_policy"] = c6.selectbox(
    "Quiz Score to Keep",
    ["keep_highest", "keep_latest"],
    index=0 if d.get("scoring_policy", "keep_highest") == "keep_highest" else 1
)

st.markdown("**Quiz Restrictions**")
d["access_code_enabled"] = st.checkbox("Require an access code", value=bool(d.get("access_code_enabled", False)))
if d["access_code_enabled"]:
    d["access_code"] = st.text_input("Access code", value=(d.get("access_code") or ""))
else:
    d["access_code"] = ""

st.markdown("**Assign / Availability (optional, ISO datetime)**")
st.caption("Example: 2026-01-20T23:59:00Z (If you don’t know, leave blank for now.)")

cc1, cc2, cc3 = st.columns(3)
d["due_at"] = cc1.text_input("Due Date (due_at)", value=(d.get("due_at") or ""))
d["unlock_at"] = cc2.text_input("Available from (unlock_at)", value=(d.get("unlock_at") or ""))
d["lock_at"] = cc3.text_input("Until (lock_at)", value=(d.get("lock_at") or ""))

st.session_state.details = d

st.divider()

# ===================================================
# QUESTIONS EDITOR + PAGING
# ===================================================
st.subheader("3) Questions")

colp1, colp2 = st.columns([1, 1])
page_size = colp1.selectbox("Questions per page", [5, 10, 15, 20, 30], index=1)
total = len(questions)
total_pages = max(1, math.ceil(total / page_size))
page = colp2.number_input("Page", min_value=1, max_value=total_pages, value=1, step=1)

start = (page - 1) * page_size
end = min(start + page_size, total)
st.caption(f"Showing questions {start+1}–{end} of {total}")

edited = [q.copy() for q in questions]

for i in range(start, end):
    q = edited[i]
    run = st.session_state.parse_run_id

    kind = (q.get("kind") or "").lower()
    preview = strip_q_prefix(q.get("question", ""))[:90]
    label_kind = "Matching" if kind == "matching" else ("Essay/Short Answer" if kind == "essay" else "MCQ")

    with st.expander(f"Q{i+1} ({label_kind}): {preview}"):
        q_text = st.text_area(
            "Question text",
            value=q.get("question", ""),
            key=f"{run}_qtext_{i}",
            height=90
        )

        q["question"] = strip_q_prefix(q_text.strip())

        if kind == "essay":
            st.info("This question will be uploaded as an ESSAY (student types the answer).")
            q["options"] = []
            q["correct"] = []
            q["multi"] = False

        elif kind == "matching":
            st.info("This question will be uploaded as MATCHING (left item → dropdown right item).")
            st.caption("Tip: if right side was a bullet list in Word, it appears joined with '; ' — that is correct.")
            pairs = q.get("pairs") or []
            new_pairs = []
            for j, p in enumerate(pairs):
                lc1, lc2 = st.columns([0.6, 0.4])
                left = lc1.text_input(
                    f"Left (row {j+1})",
                    value=p.get("left",""),
                    key=f"{run}_match_{i}_l_{j}"
                )
                right = lc2.text_input(
                    f"Right (row {j+1})",
                    value=p.get("right",""),
                    key=f"{run}_match_{i}_r_{j}"
                )

                if left.strip() and right.strip():
                    new_pairs.append({"left": left.strip(), "right": right.strip()})
            q["pairs"] = new_pairs

        else:
            opts = q.get("options", []) or []
            correct_set = set(q.get("correct", []) or [])

            st.write("**Options** (tick ✅ for correct answer)")
            desired_n = st.number_input(
                "Number of options",
                min_value=2,
                max_value=60,
                value=max(2, len(opts)),
                step=1,
                key=f"{run}_q{i}_nopts",
            )
            if desired_n > len(opts):
                opts = opts + ([""] * (desired_n - len(opts)))
            elif desired_n < len(opts):
                opts = opts[:desired_n]
                correct_set = {x for x in correct_set if x < desired_n}

            if any((o or "").startswith("⚠ Option text not extracted") for o in opts):
                st.info("This question's options were not extracted (often because they are images/shapes). Paste the options below to replace them quickly.")
                pasted = st.text_area(
                    "Paste options (one per line)",
                    value="",
                    key=f"{run}_q{i}_pasted_opts",
                    height=120,
                    placeholder="Example:\nDescription of food item\nDate and time produced\nUse by or best before date",
                )
                new_from_paste = parse_pasted_options(pasted)
                if len(new_from_paste) >= 2:
                    opts = new_from_paste
                    correct_set = set()

            new_opts = []
            new_correct = []

            for j, opt in enumerate(opts):
                oc1, oc2 = st.columns([0.12, 0.88])
                is_corr = oc1.checkbox(
                    "",
                    value=(j in correct_set),
                    key=f"{run}_q{i}_corr_{j}"
                )
                opt_text = oc2.text_input(
                    f"Option {j+1}",
                    value=opt,
                    key=f"{run}_q{i}_opt_{j}"
                )

                new_opts.append(opt_text.strip())
                if is_corr:
                    new_correct.append(j)

            add_opt = st.text_input(
                "New option text (optional)",
                value="",
                key=f"{run}_q{i}_newopt"
            )

            if add_opt.strip():
                new_opts.append(add_opt.strip())

            cleaned_opts = []
            idx_map = {}
            for old_index, txt in enumerate(new_opts):
                if txt.strip():
                    idx_map[old_index] = len(cleaned_opts)
                    cleaned_opts.append(txt.strip())

            remapped_correct = []
            for old_i in new_correct:
                if old_i in idx_map:
                    remapped_correct.append(idx_map[old_i])

            q["options"] = cleaned_opts
            q["correct"] = sorted(set(remapped_correct))

            qlower = (q.get("question") or "").lower()
            q["multi"] = (
                ("select" in qlower and any(w in qlower for w in ["two", "three", "four", "five"]))
                or ("apply" in qlower)
                or (len(q["correct"]) > 1)
            )

st.session_state.questions = edited

st.divider()

# ===================================================
# SAVE ACTIONS
# ===================================================
st.subheader("4) Save to Canvas")

colS1, colS2 = st.columns([1, 1])
save_draft = colS1.button("💾 Save to Canvas (Draft)")
save_publish = colS2.button("🚀 Save & Publish")

def validate_before_upload(qs: list[dict]) -> list[str]:
    problems = []
    for i, q in enumerate(qs, start=1):
        kind = (q.get("kind") or "").lower()
        qt = (q.get("question") or "").strip()

        if len(qt) < 10:
            problems.append(f"Q{i}: question text too short.")

        if kind == "matching":
            pairs = q.get("pairs") or []
            if len(pairs) < 2:
                problems.append(f"Q{i}: matching needs at least 2 pairs.")
            continue

        if kind == "essay":
            continue

        opts = q.get("options") or []
        corr = q.get("correct") or []
        if len(opts) >= 2 and len(corr) == 0:
            problems.append(f"Q{i}: no correct answer selected (red not detected or tick ✅).")
    return problems

if save_draft or save_publish:
    qs = st.session_state.questions
    probs = validate_before_upload(qs)
    if probs:
        st.error("Please fix these issues before uploading:")
        for p in probs[:15]:
            st.write(f"- {p}")
        st.stop()

    base_title = (quiz_title or "").strip() or default_title
    try:
        existing_titles = get_existing_quiz_titles(canvas_base_url, course_id, canvas_token)
        final_title = generate_unique_title(base_title, existing_titles)

        with st.spinner("Creating quiz in Canvas..."):
            quiz_id = create_canvas_quiz(
                canvas_base_url=canvas_base_url,
                course_id=course_id,
                canvas_token=canvas_token,
                title=final_title,
                description_html=quiz_instructions,
                settings=st.session_state.details
            )

        with st.spinner("Uploading questions..."):
            for q in qs:
                add_question_to_quiz(canvas_base_url, course_id, canvas_token, quiz_id, q)

        if save_publish:
            with st.spinner("Publishing quiz..."):
                publish_quiz(canvas_base_url, course_id, canvas_token, quiz_id)

        st.success("✅ Done!")
        st.write(f"**Quiz title:** {final_title}")
        st.write(f"**Quiz ID:** {quiz_id}")
        st.write(f"**Course ID:** {course_id}")
        st.info("Quiz published ✅" if save_publish else "Quiz saved as draft (unpublished).")

    except Exception as e:
        st.error(f"❌ Upload failed: {e}")

st.caption("Token login only (Canvas API does not support username/password).")
